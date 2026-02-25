"""
Tests for the improved parse_docx module.

Covers:
- Heading-style detection (existing behaviour)
- Heuristic fallback with bold/uppercase normal paragraphs
- Section markers inside Word table cells
- Debug sections helper
"""

import tempfile

import pytest
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from carbongpt.tools.parse_docx import parse_docx, debug_sections, _is_section_marker


class TestSectionMarkerRegex:
    def test_key_project_information(self):
        assert _is_section_marker("KEY PROJECT INFORMATION")

    def test_key_project_information_lowercase(self):
        assert _is_section_marker("Key Project Information")

    def test_section_a(self):
        assert _is_section_marker("SECTION A")

    def test_section_a_with_title(self):
        assert _is_section_marker("SECTION A. DESCRIPTION OF PROJECT")

    def test_section_b_dash(self):
        assert _is_section_marker("SECTION B – IMPLEMENTATION OF PROJECT")

    def test_section_g(self):
        assert _is_section_marker("SECTION G. STAKEHOLDER INPUTS AND LEGAL DISPUTES")

    def test_not_section(self):
        assert not _is_section_marker("This is normal text")

    def test_not_section_partial(self):
        assert not _is_section_marker("See SECTION A for details")


class TestHeadingStylePass:
    @staticmethod
    def _make_heading_docx(sections: dict[str, str]) -> str:
        doc = Document()
        for heading, body in sections.items():
            doc.add_heading(heading, level=1)
            doc.add_paragraph(body)
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        return tmp.name

    def test_heading_styles_detected(self):
        path = self._make_heading_docx({
            "SECTION A": "Description content.",
            "SECTION B": "Implementation content.",
        })
        result = parse_docx(path)
        assert result["debug"]["method"] == "heading_styles"
        assert "SECTION A" in result["sections"]
        assert "SECTION B" in result["sections"]


class TestHeuristicFallback:
    @staticmethod
    def _make_plain_docx(section_entries: list[tuple[str, str]]) -> str:
        """Create a docx with section titles as bold normal paragraphs (no Heading style)."""
        doc = Document()
        for title, body in section_entries:
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(14)
            doc.add_paragraph(body)
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        return tmp.name

    def test_bold_uppercase_sections_detected(self):
        path = self._make_plain_docx([
            ("KEY PROJECT INFORMATION", "GS ID: GS-1234 Title of the project: Test"),
            ("SECTION A. DESCRIPTION OF PROJECT", "Project description text."),
            ("SECTION B. IMPLEMENTATION OF PROJECT", "Implementation details."),
            ("SECTION C. DESCRIPTION OF MONITORING SYSTEM APPLIED BY THE PROJECT", "Monitoring system."),
            ("SECTION D. DATA AND PARAMETERS", "Parameters measured."),
            ("SECTION E. CALCULATION OF SDG IMPACTS", "Emission reductions."),
            ("SECTION F. SAFEGUARDS REPORTING", "Safeguards applied."),
            ("SECTION G. STAKEHOLDER INPUTS AND LEGAL DISPUTES", "No disputes."),
        ])

        result = parse_docx(path)
        assert result["debug"]["method"] == "heuristic"
        section_names = list(result["sections"].keys())

        assert any("KEY PROJECT INFORMATION" in s for s in section_names)
        for letter in "ABCDEFG":
            assert any(f"SECTION {letter}" in s for s in section_names), f"SECTION {letter} not found"

    def test_preamble_before_sections(self):
        path = self._make_plain_docx([
            ("Monitoring Report", "Publication date 2024."),
            ("KEY PROJECT INFORMATION", "GS ID: GS-5678"),
            ("SECTION A. DESCRIPTION OF PROJECT", "Description."),
        ])

        result = parse_docx(path)
        section_names = list(result["sections"].keys())
        assert any("KEY PROJECT INFORMATION" in s for s in section_names)
        assert any("SECTION A" in s for s in section_names)

    def test_section_body_text_accumulated(self):
        path = self._make_plain_docx([
            ("KEY PROJECT INFORMATION", "GS ID: GS-1234"),
            ("SECTION A. DESCRIPTION OF PROJECT", "First paragraph."),
        ])

        result = parse_docx(path)
        kpi_key = [k for k in result["sections"] if "KEY PROJECT INFORMATION" in k][0]
        assert "GS-1234" in result["sections"][kpi_key]


class TestTableTextDetection:
    @staticmethod
    def _make_table_docx() -> str:
        """Create a docx where KEY PROJECT INFORMATION is inside a table cell."""
        doc = Document()
        doc.add_paragraph("Monitoring Report")

        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "KEY PROJECT INFORMATION"
        table.cell(0, 1).text = ""
        table.cell(1, 0).text = "GS ID"
        table.cell(1, 1).text = "GS-9999"

        p = doc.add_paragraph()
        run = p.add_run("SECTION A. DESCRIPTION OF PROJECT")
        run.bold = True
        doc.add_paragraph("Project description goes here.")

        p2 = doc.add_paragraph()
        run2 = p2.add_run("SECTION B. IMPLEMENTATION OF PROJECT")
        run2.bold = True
        doc.add_paragraph("Implementation details.")

        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        return tmp.name

    def test_kpi_in_table_detected(self):
        path = self._make_table_docx()
        result = parse_docx(path)
        section_names = list(result["sections"].keys())

        assert any("KEY PROJECT INFORMATION" in s for s in section_names), (
            f"KEY PROJECT INFORMATION not found in {section_names}"
        )

    def test_sections_after_table_detected(self):
        path = self._make_table_docx()
        result = parse_docx(path)
        section_names = list(result["sections"].keys())
        assert any("SECTION A" in s for s in section_names)
        assert any("SECTION B" in s for s in section_names)


class TestDebugSections:
    @staticmethod
    def _make_simple_docx() -> str:
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("KEY PROJECT INFORMATION").bold = True
        doc.add_paragraph("GS ID: GS-1234")
        p2 = doc.add_paragraph()
        p2.add_run("SECTION A. DESCRIPTION OF PROJECT").bold = True
        doc.add_paragraph("Description.")
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        return tmp.name

    def test_debug_returns_expected_keys(self):
        path = self._make_simple_docx()
        info = debug_sections(path)

        assert "raw_paragraphs_first_30" in info
        assert "section_markers" in info
        assert "final_section_names" in info
        assert "detection_method" in info
        assert "total_elements" in info

    def test_debug_markers_found(self):
        path = self._make_simple_docx()
        info = debug_sections(path)
        marker_texts = [m["text"] for m in info["section_markers"]]
        assert any("KEY PROJECT INFORMATION" in t for t in marker_texts)
        assert any("SECTION A" in t for t in marker_texts)


class TestRealTemplateDocx:
    def test_internal_template_parses(self):
        from pathlib import Path
        template_path = Path(__file__).resolve().parent.parent / "templates" / "goldstandard" / "MR_v1_1.docx"
        if not template_path.exists():
            pytest.skip("Template docx not available")

        result = parse_docx(str(template_path))
        section_names = list(result["sections"].keys())

        assert any("KEY PROJECT INFORMATION" in s.upper() for s in section_names), (
            f"KEY PROJECT INFORMATION not found in template. Got: {section_names}"
        )
        assert len(section_names) >= 3, f"Expected at least 3 sections, got {section_names}"
