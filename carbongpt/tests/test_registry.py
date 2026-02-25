"""
Tests for the template registry and /analyze-selected endpoint.
"""

import tempfile

import pytest
from docx import Document

from carbongpt.templates.registry import list_standards, list_doc_types, list_versions, lookup
from carbongpt.core.orchestrator import run_selected_analysis


class TestRegistryLookup:
    def test_list_standards(self):
        standards = list_standards()
        assert "GoldStandard" in standards

    def test_list_doc_types(self):
        doc_types = list_doc_types("GoldStandard")
        assert "MR" in doc_types
        assert "PDD" in doc_types

    def test_list_versions_mr(self):
        versions = list_versions("GoldStandard", "MR")
        assert "MR_v1_1" in versions

    def test_lookup_valid(self):
        entry = lookup("GoldStandard", "MR", "MR_v1_1")
        assert entry is not None
        assert "template_path" in entry
        assert "rules_path" in entry
        assert entry["template_path"].endswith(".docx")
        assert entry["rules_path"].endswith(".yaml")

    def test_lookup_invalid(self):
        assert lookup("NonExistent", "MR", "MR_v1_1") is None
        assert lookup("GoldStandard", "MR", "DOES_NOT_EXIST") is None

    def test_lookup_pdd_placeholder(self):
        entry = lookup("GoldStandard", "PDD", "PDD_v1_0")
        assert entry is not None
        assert entry["template_path"] is None
        assert entry["rules_path"] is None


class TestAnalyzeSelected:
    @staticmethod
    def _make_docx(sections: dict[str, str]) -> str:
        doc = Document()
        for heading, body in sections.items():
            doc.add_heading(heading, level=1)
            doc.add_paragraph(body)
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        return tmp.name

    def test_valid_analysis(self):
        path = self._make_docx({
            "KEY PROJECT INFORMATION": (
                "GS ID: GS-1234 Title of the project: Solar Cookstoves "
                "Version number of the monitoring report: v2.0 "
                "Completion date: 15/06/2024 "
                "Monitoring period number: MP #3 "
                "Duration of this monitoring period: from 01/01/2024 to 31/12/2024 "
                "Host Country: Kenya "
                "Methodology applied: AMS-II.G version 09"
            ),
            "SECTION A. DESCRIPTION OF PROJECT": "Solar cookstove distribution.",
            "SECTION B. IMPLEMENTATION OF PROJECT": "Implemented in rural Kenya.",
            "SECTION C. DESCRIPTION OF MONITORING SYSTEM APPLIED BY THE PROJECT": "Monitoring system details.",
            "SECTION D. DATA AND PARAMETERS": "Parameters measured.",
            "SECTION E. CALCULATION OF SDG IMPACTS": (
                "Baseline emissions: 15000 tCO2e. "
                "Project emissions: 500 tCO2e. "
                "Emission reductions: 14500 tCO2e."
            ),
            "SECTION F. SAFEGUARDS REPORTING": "Environmental safeguards applied comprehensively for the project area.",
            "SECTION G. STAKEHOLDER INPUTS AND LEGAL DISPUTES": "No disputes reported during monitoring period for the project.",
        })

        result = run_selected_analysis(
            standard="GoldStandard",
            doc_type="MR",
            version="MR_v1_1",
            user_doc_path=path,
        )

        assert result.standard == "GoldStandard"
        assert result.doc_type == "MR"
        assert result.version == "MR_v1_1"
        assert result.compliance_score <= 100
        assert result.compliance_score >= 0
        assert len(result.sections_found) > 0
        assert len(result.template_sections) > 0

    def test_invalid_registry_entry(self):
        path = self._make_docx({"Test": "content"})
        with pytest.raises(ValueError, match="No template registered"):
            run_selected_analysis(
                standard="NonExistent",
                doc_type="MR",
                version="MR_v1_1",
                user_doc_path=path,
            )

    def test_pdd_placeholder_raises(self):
        path = self._make_docx({"Test": "content"})
        with pytest.raises(ValueError, match="not yet available"):
            run_selected_analysis(
                standard="GoldStandard",
                doc_type="PDD",
                version="PDD_v1_0",
                user_doc_path=path,
            )

    def test_missing_sections_produce_findings(self):
        path = self._make_docx({
            "KEY PROJECT INFORMATION": "GS ID: GS-9999 Title of the project: Test",
        })

        result = run_selected_analysis(
            standard="GoldStandard",
            doc_type="MR",
            version="MR_v1_1",
            user_doc_path=path,
        )

        assert result.compliant is False
        assert result.compliance_score < 100
        assert len(result.findings) > 0
        error_msgs = [f.message for f in result.findings if f.severity == "ERROR"]
        assert len(error_msgs) > 0
