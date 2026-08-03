"""
Tests for carbongpt.core.doc_exporter's checkbox-ticking (docs/SPEC-05.md
T6). Runs against the real VPA-DD v2.3/v3.0 .docx files already tracked
in document_repository/ — no network, no database.

Finding that motivated this fix (03.08.2026): _check_checkbox_in_cell
only recognised the legacy Word form-field mechanism
(w:ffData/w:checkBox, Wingdings w:sym char codes). Checked directly
against both templates: v2.3 (the guide in production use before this
session) has ZERO legacy checkboxes and 495 modern structured-document-tag
checkboxes (w:sdt/w14:checkbox) ; v3.0 has 479, also all modern. The
legacy path was effectively dead code — checkbox-ticking has silently
done nothing for either template version, not just a v3.0-readiness gap.
"""
from pathlib import Path

import docx
import pytest

from carbongpt.core.doc_exporter import _check_checkbox_in_cell

REPO_DIR = Path(__file__).parent.parent.parent / "document_repository"
V2_3_PATH = REPO_DIR / "5629069d7319be57_T-PreReview_v2.3-VPA-Design-Document.docx"
V3_0_PATH = REPO_DIR / "85924c132bed502d_T-PAA_PreReview_V3.0_VPA-Design-Document.docx"

NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
W14_NS = "{http://schemas.microsoft.com/office/word/2010/wordml}"

requires_v2_3 = pytest.mark.skipif(not V2_3_PATH.exists(), reason="VPA-DD v2.3 not ingested yet")
requires_v3_0 = pytest.mark.skipif(not V3_0_PATH.exists(), reason="VPA-DD v3.0 not ingested yet")


def _checkbox_states(cell):
    """[(checked_val, glyph_char), ...] for every w14 checkbox in a cell."""
    states = []
    for sdt in cell._tc.findall(".//w:sdt", NS):
        cb = sdt.find(f".//{W14_NS}checkbox")
        if cb is None:
            continue
        checked = cb.find(f"{W14_NS}checked")
        content = sdt.find("w:sdtContent", NS)
        glyph = None
        if content is not None:
            for t_el in content.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                if t_el.text and ord(t_el.text[:1] or "\0") in (0x2610, 0x2612):
                    glyph = t_el.text
        states.append((checked.get(f"{W14_NS}val") if checked is not None else None, glyph))
    return states


@requires_v3_0
class TestModernCheckboxTickingV3_0:
    def test_no_legacy_checkboxes_exist_in_v3_0(self):
        # Confirms the premise: the legacy path alone would tick nothing.
        doc = docx.Document(str(V3_0_PATH))
        assert len(doc.element.body.findall(".//w:ffData", NS)) == 0

    def test_ticking_a_real_checkbox_flips_state_and_glyph(self):
        doc = docx.Document(str(V3_0_PATH))
        cell = doc.tables[1].rows[0].cells[1]  # "Real-case / Regular" — confirmed 03.08.2026
        before = _checkbox_states(cell)
        assert before, "expected at least one w14 checkbox in this cell"
        assert all(val == "0" and glyph == "☐" for val, glyph in before)

        _check_checkbox_in_cell(cell._element, "Real-case", NS)

        after = _checkbox_states(cell)
        assert all(val == "1" and glyph == "☒" for val, glyph in after)

    def test_non_matching_text_ticks_nothing(self):
        doc = docx.Document(str(V3_0_PATH))
        cell = doc.tables[1].rows[0].cells[1]
        _check_checkbox_in_cell(cell._element, "Some Unrelated Label Text", NS)
        after = _checkbox_states(cell)
        assert all(val == "0" and glyph == "☐" for val, glyph in after)


@requires_v2_3
class TestModernCheckboxTickingV2_3:
    def test_no_legacy_checkboxes_exist_in_v2_3_either(self):
        # The "already broken before this session" finding, not v3.0-specific.
        doc = docx.Document(str(V2_3_PATH))
        assert len(doc.element.body.findall(".//w:ffData", NS)) == 0
        w14_count = sum(
            1 for sdt in doc.element.body.findall(".//w:sdt", NS)
            if sdt.find(f".//{W14_NS}checkbox") is not None
        )
        assert w14_count > 0

    def test_ticking_a_real_checkbox_flips_only_the_matching_paragraph(self):
        # Unlike v3.0's "Real-case / Regular" (same paragraph, both tick —
        # see the other test class), v2.3 puts "Real case VPA" and
        # "Regular VPA" in TWO SEPARATE paragraphs (confirmed 03.08.2026) —
        # matching is done per-paragraph, so only the matched one should tick.
        doc = docx.Document(str(V2_3_PATH))
        cell = doc.tables[0].rows[0].cells[1]  # "Real case VPA" / "Regular VPA"
        before = _checkbox_states(cell)
        assert len(before) == 2
        assert all(val == "0" for val, _glyph in before)

        _check_checkbox_in_cell(cell._element, "Real case VPA", NS)

        after = _checkbox_states(cell)
        assert after[0] == ("1", "☒")
        assert after[1] == ("0", "☐")
