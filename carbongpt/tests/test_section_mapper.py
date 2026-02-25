"""
Tests for the section_mapper and its integration with the rule engine.

Covers:
- exact match
- fuzzy match (e.g. "B.1 Monitoring period" matching "Monitoring Period")
- missing section detection
- normalize_heading edge cases
"""

import tempfile
from pathlib import Path

import pytest
from docx import Document

from carbongpt.tools.section_mapper import map_sections, normalize_heading
from carbongpt.tools.rule_engine import run_template_rules
from carbongpt.tools.parse_docx import parse_docx


# ---------------------------------------------------------------------------
# normalize_heading
# ---------------------------------------------------------------------------

class TestNormalizeHeading:
    def test_lowercase_and_strip(self):
        assert normalize_heading("  Project Description  ") == "project description"

    def test_remove_punctuation(self):
        assert normalize_heading("B.1 Monitoring Period") == "b1 monitoring period"

    def test_collapse_spaces(self):
        assert normalize_heading("Data   Quality   Management") == "data quality management"

    def test_mixed(self):
        assert normalize_heading("  A.2.1 -- Emission Reductions!!  ") == "a21 emission reductions"

    def test_empty_string(self):
        assert normalize_heading("") == ""


# ---------------------------------------------------------------------------
# map_sections — exact match
# ---------------------------------------------------------------------------

class TestMapSectionsExact:
    def test_exact_match_all_present(self):
        expected = ["Monitoring Period", "Project Description"]
        found = ["Monitoring Period", "Project Description", "Extra Section"]
        result = map_sections(expected, found)
        assert result["Monitoring Period"] == "Monitoring Period"
        assert result["Project Description"] == "Project Description"

    def test_exact_match_case_insensitive(self):
        expected = ["monitoring period"]
        found = ["Monitoring Period"]
        result = map_sections(expected, found)
        assert result["monitoring period"] == "Monitoring Period"


# ---------------------------------------------------------------------------
# map_sections — fuzzy match
# ---------------------------------------------------------------------------

class TestMapSectionsFuzzy:
    def test_numbered_prefix_matches(self):
        expected = ["Monitoring Period"]
        found = ["B.1 Monitoring Period"]
        result = map_sections(expected, found, threshold=70)
        assert result["Monitoring Period"] == "B.1 Monitoring Period"

    def test_slight_wording_variation(self):
        expected = ["Project Description"]
        found = ["Description of the Project"]
        result = map_sections(expected, found, threshold=60)
        assert result["Project Description"] == "Description of the Project"

    def test_below_threshold_returns_none(self):
        expected = ["Monitoring Period"]
        found = ["Completely Unrelated Heading"]
        result = map_sections(expected, found, threshold=85)
        assert result["Monitoring Period"] is None


# ---------------------------------------------------------------------------
# map_sections — missing section
# ---------------------------------------------------------------------------

class TestMapSectionsMissing:
    def test_no_found_headings(self):
        expected = ["Monitoring Period", "Project Description"]
        result = map_sections(expected, [], threshold=85)
        assert result["Monitoring Period"] is None
        assert result["Project Description"] is None

    def test_partial_match(self):
        expected = ["Monitoring Period", "Baseline Emissions", "Safeguards"]
        found = ["Monitoring Period"]
        result = map_sections(expected, found, threshold=85)
        assert result["Monitoring Period"] == "Monitoring Period"
        assert result["Baseline Emissions"] is None
        assert result["Safeguards"] is None


# ---------------------------------------------------------------------------
# run_template_rules integration
# ---------------------------------------------------------------------------

class TestRunTemplateRules:
    def test_all_present_no_findings(self):
        expected = ["Monitoring Period", "Project Description"]
        sections = {"Monitoring Period": "text", "Project Description": "text"}
        findings, section_map = run_template_rules(expected, sections)
        assert len(findings) == 0
        assert section_map["Monitoring Period"] == "Monitoring Period"

    def test_missing_section_produces_finding(self):
        expected = ["Monitoring Period", "Project Description"]
        sections = {"Monitoring Period": "text"}
        findings, section_map = run_template_rules(expected, sections)
        assert len(findings) == 1
        assert findings[0].message == "Missing required section: Project Description"
        assert findings[0].severity == "ERROR"
        assert section_map["Project Description"] is None

    def test_fuzzy_match_passes(self):
        expected = ["Monitoring Period"]
        sections = {"B.1 Monitoring Period": "text"}
        findings, section_map = run_template_rules(expected, sections, threshold=70)
        assert len(findings) == 0
        assert section_map["Monitoring Period"] == "B.1 Monitoring Period"


# ---------------------------------------------------------------------------
# End-to-end with real docx files
# ---------------------------------------------------------------------------

class TestEndToEndDocx:
    @staticmethod
    def _make_docx(headings: list[str]) -> str:
        doc = Document()
        for h in headings:
            doc.add_heading(h, level=1)
            doc.add_paragraph(f"Content for {h}.")
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        return tmp.name

    def test_template_vs_user_fuzzy(self):
        template_path = self._make_docx(["Monitoring Period", "Project Description"])
        user_path = self._make_docx(["B.1 Monitoring Period", "A.1 Project Description"])

        template_sections = list(parse_docx(template_path)["sections"].keys())
        user_sections = parse_docx(user_path)["sections"]

        findings, section_map = run_template_rules(
            template_sections, user_sections, threshold=70
        )

        assert len(findings) == 0
        assert section_map["Monitoring Period"] == "B.1 Monitoring Period"
        assert section_map["Project Description"] == "A.1 Project Description"

    def test_template_vs_user_missing(self):
        template_path = self._make_docx(
            ["Monitoring Period", "Project Description", "Safeguards"]
        )
        user_path = self._make_docx(["Monitoring Period"])

        template_sections = list(parse_docx(template_path)["sections"].keys())
        user_sections = parse_docx(user_path)["sections"]

        findings, section_map = run_template_rules(
            template_sections, user_sections, threshold=85
        )

        assert len(findings) == 2
        missing_names = {f.message.replace("Missing required section: ", "") for f in findings}
        assert "Project Description" in missing_names
        assert "Safeguards" in missing_names
