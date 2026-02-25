"""
Tests for all rule types and compliance_score computation.

Covers:
- required_field: found, missing, silently skipped when section missing
- date_format_ddmmyyyy: correct format, wrong format, no dates, section missing
- not_applicable_required_when_blank: short with N/A, short without, long text, section missing
- compliance_score calculation
- text normalization
- end-to-end with real docx files
"""

import tempfile

import pytest
from docx import Document

from carbongpt.core.models import Finding, compute_status
from carbongpt.tools.rule_engine import (
    compute_compliance_score,
    _check_required_field,
    _check_date_format_ddmmyyyy,
    _check_not_applicable_required_when_blank,
    _check_must_mention_keywords,
    _normalize_text,
)
from carbongpt.tools.section_mapper import map_sections
from carbongpt.tools.regex_utils import any_pattern_matches, find_all_matches, is_ddmmyyyy
from carbongpt.tools.parse_docx import parse_docx
from carbongpt.core.orchestrator import run_analysis


# ---------------------------------------------------------------------------
# regex_utils
# ---------------------------------------------------------------------------

class TestRegexUtils:
    def test_any_pattern_matches_date(self):
        assert any_pattern_matches("Started on 2024-01-15", [r"\d{4}-\d{2}-\d{2}"])

    def test_any_pattern_matches_none(self):
        assert not any_pattern_matches("No dates here", [r"\d{4}-\d{2}-\d{2}"])

    def test_find_all_matches(self):
        text = "Dates: 2024-01-01 and 15/06/2024"
        matches = find_all_matches(text, [r"\d{4}-\d{2}-\d{2}", r"\d{2}/\d{2}/\d{4}"])
        assert "2024-01-01" in matches
        assert "15/06/2024" in matches

    def test_is_ddmmyyyy_valid(self):
        assert is_ddmmyyyy("15/06/2024")
        assert is_ddmmyyyy("01/01/2023")

    def test_is_ddmmyyyy_invalid(self):
        assert not is_ddmmyyyy("2024-01-15")
        assert not is_ddmmyyyy("15-06-2024")
        assert not is_ddmmyyyy("June 15, 2024")


# ---------------------------------------------------------------------------
# _normalize_text
# ---------------------------------------------------------------------------

class TestNormalizeText:
    def test_collapses_whitespace(self):
        assert _normalize_text("hello   world") == "hello world"

    def test_replaces_newlines(self):
        assert _normalize_text("line1\nline2\r\nline3") == "line1 line2 line3"

    def test_strips(self):
        assert _normalize_text("  hello  ") == "hello"

    def test_complex(self):
        text = "  GS ID:\n  GS-1234  \n  Title of\r\n  the project  "
        result = _normalize_text(text)
        assert "GS ID:" in result
        assert "GS-1234" in result
        assert "\n" not in result


# ---------------------------------------------------------------------------
# _check_required_field — field found
# ---------------------------------------------------------------------------

class TestRequiredFieldFound:
    def test_date_present(self):
        rule = {
            "id": "F001", "type": "required_field",
            "section": "KEY PROJECT INFORMATION", "field_name": "start",
            "severity": "ERROR", "patterns": [r"\d{4}-\d{2}-\d{2}"],
        }
        sections = {"KEY PROJECT INFORMATION": "Started 2024-01-01."}
        section_map = {"KEY PROJECT INFORMATION": "KEY PROJECT INFORMATION"}
        assert _check_required_field(rule, sections, section_map) is None

    def test_keyword_present(self):
        rule = {
            "id": "F004", "type": "required_field",
            "section": "KEY PROJECT INFORMATION", "field_name": "completion_date",
            "severity": "ERROR", "patterns": [r"(?i)completion\s+date"],
        }
        sections = {"KEY PROJECT INFORMATION": "Completion date of the monitoring report: 15/06/2024"}
        section_map = {"KEY PROJECT INFORMATION": "KEY PROJECT INFORMATION"}
        assert _check_required_field(rule, sections, section_map) is None

    def test_fuzzy_matched_heading(self):
        sections = {"B.1 Monitoring Period": "Data from 2024-06-01."}
        section_map = map_sections(["Monitoring Period"], list(sections.keys()), threshold=70)
        rule = {
            "id": "F002", "type": "required_field",
            "section": "Monitoring Period", "field_name": "start",
            "severity": "ERROR", "patterns": [r"\d{4}-\d{2}-\d{2}"],
        }
        assert _check_required_field(rule, sections, section_map) is None

    def test_gs_template_field_patterns(self):
        text = (
            "Title of the project (s) covered by monitoring report "
            "Version number of the monitoring report "
            "Completion date of the monitoring report "
            "Monitoring period number "
            "Duration of this monitoring period "
            "Host Country "
            "Methodology (ies) applied and version number"
        )
        section_map = {"KEY PROJECT INFORMATION": "KEY PROJECT INFORMATION"}
        sections = {"KEY PROJECT INFORMATION": text}

        patterns_map = {
            "title_of_project": [r"(?i)title\s+of\s+(?:the\s+)?project"],
            "version_number": [r"(?i)version\s+number\s+of\s+the\s+monitoring\s+report"],
            "completion_date": [r"(?i)completion\s+date"],
            "monitoring_period_number": [r"(?i)monitoring\s+period\s+number"],
            "duration": [r"(?i)duration\s+of\s+(?:this\s+)?(?:monitoring\s+)?period"],
            "host_country": [r"(?i)host\s+country"],
            "methodology": [r"(?i)methodolog(?:y|ies)\s+.*version"],
        }

        for field_name, patterns in patterns_map.items():
            rule = {
                "id": f"F_{field_name}", "type": "required_field",
                "section": "KEY PROJECT INFORMATION", "field_name": field_name,
                "severity": "ERROR", "patterns": patterns,
            }
            result = _check_required_field(rule, sections, section_map)
            assert result is None, f"Pattern for {field_name} should match but didn't"


# ---------------------------------------------------------------------------
# _check_required_field — field missing
# ---------------------------------------------------------------------------

class TestRequiredFieldMissing:
    def test_no_pattern_matches(self):
        rule = {
            "id": "F001", "type": "required_field",
            "section": "KEY PROJECT INFORMATION", "field_name": "gs_id",
            "severity": "ERROR", "patterns": [r"(?i)GS\s*[-:]?\s*\d{3,}"],
        }
        sections = {"KEY PROJECT INFORMATION": "No project ID here."}
        section_map = {"KEY PROJECT INFORMATION": "KEY PROJECT INFORMATION"}
        finding = _check_required_field(rule, sections, section_map)
        assert finding is not None
        assert "gs_id" in finding.message
        assert finding.severity == "ERROR"


# ---------------------------------------------------------------------------
# _check_required_field — section missing => silently skipped
# ---------------------------------------------------------------------------

class TestRequiredFieldSectionMissing:
    def test_returns_none_when_section_missing(self):
        rule = {
            "id": "F001", "type": "required_field",
            "section": "KEY PROJECT INFORMATION", "field_name": "gs_id",
            "severity": "ERROR", "patterns": [r"(?i)GS\s+ID"],
        }
        sections = {"SECTION A": "Some text."}
        section_map = {"KEY PROJECT INFORMATION": None}
        assert _check_required_field(rule, sections, section_map) is None

    def test_returns_none_when_no_sections_at_all(self):
        rule = {
            "id": "F003", "type": "required_field",
            "section": "SECTION F", "field_name": "safeguards",
            "severity": "WARNING", "patterns": [r"(?i)safeguard"],
        }
        assert _check_required_field(rule, {}, {"SECTION F": None}) is None


# ---------------------------------------------------------------------------
# _check_required_field — text aggregation across duplicate sections
# ---------------------------------------------------------------------------

class TestRequiredFieldTextAggregation:
    def test_matches_across_duplicate_section_names(self):
        rule = {
            "id": "F001", "type": "required_field",
            "section": "KEY PROJECT INFORMATION", "field_name": "host_country",
            "severity": "ERROR", "patterns": [r"(?i)host\s+country"],
        }
        sections = {
            "Key Project Information": "Host Country: Kenya",
            "KEY PROJECT INFORMATION": "Some other text here.",
        }
        section_map = {"KEY PROJECT INFORMATION": "KEY PROJECT INFORMATION"}
        assert _check_required_field(rule, sections, section_map) is None


# ---------------------------------------------------------------------------
# _check_date_format_ddmmyyyy
# ---------------------------------------------------------------------------

class TestDateFormatDDMMYYYY:
    def _rule(self, section="KEY PROJECT INFORMATION"):
        return {
            "id": "D001", "type": "date_format_ddmmyyyy",
            "section": section, "severity": "WARNING",
            "date_patterns": [r"\d{4}-\d{2}-\d{2}", r"\d{2}-\d{2}-\d{4}", r"\d{2}/\d{2}/\d{4}"],
        }

    def test_correct_format_no_finding(self):
        sections = {"KEY PROJECT INFORMATION": "Completion date: 15/06/2024"}
        section_map = {"KEY PROJECT INFORMATION": "KEY PROJECT INFORMATION"}
        assert _check_date_format_ddmmyyyy(self._rule(), sections, section_map) is None

    def test_wrong_format_yyyy_mm_dd(self):
        sections = {"KEY PROJECT INFORMATION": "Completion date: 2024-06-15"}
        section_map = {"KEY PROJECT INFORMATION": "KEY PROJECT INFORMATION"}
        finding = _check_date_format_ddmmyyyy(self._rule(), sections, section_map)
        assert finding is not None
        assert "2024-06-15" in finding.message
        assert finding.severity == "WARNING"

    def test_wrong_format_dd_dash_mm_yyyy(self):
        sections = {"KEY PROJECT INFORMATION": "Date: 15-06-2024"}
        section_map = {"KEY PROJECT INFORMATION": "KEY PROJECT INFORMATION"}
        finding = _check_date_format_ddmmyyyy(self._rule(), sections, section_map)
        assert finding is not None
        assert "15-06-2024" in finding.message

    def test_mixed_formats_reports_bad_only(self):
        sections = {"KEY PROJECT INFORMATION": "From 01/01/2024 to 2024-12-31"}
        section_map = {"KEY PROJECT INFORMATION": "KEY PROJECT INFORMATION"}
        finding = _check_date_format_ddmmyyyy(self._rule(), sections, section_map)
        assert finding is not None
        assert "2024-12-31" in finding.message
        assert "01/01/2024" not in finding.message

    def test_no_dates_no_finding(self):
        sections = {"KEY PROJECT INFORMATION": "No date information here."}
        section_map = {"KEY PROJECT INFORMATION": "KEY PROJECT INFORMATION"}
        assert _check_date_format_ddmmyyyy(self._rule(), sections, section_map) is None

    def test_section_missing_no_finding(self):
        sections = {}
        section_map = {"KEY PROJECT INFORMATION": None}
        assert _check_date_format_ddmmyyyy(self._rule(), sections, section_map) is None


# ---------------------------------------------------------------------------
# _check_not_applicable_required_when_blank
# ---------------------------------------------------------------------------

class TestNotApplicableRequiredWhenBlank:
    def _rule(self, section="SECTION F", min_chars=40):
        return {
            "id": "NA001", "type": "not_applicable_required_when_blank",
            "section": section, "severity": "WARNING", "min_chars": min_chars,
        }

    def test_long_text_no_finding(self):
        text = "This section provides comprehensive safeguards reporting and assessment for the project."
        sections = {"SECTION F": text}
        section_map = {"SECTION F": "SECTION F"}
        assert _check_not_applicable_required_when_blank(self._rule(), sections, section_map) is None

    def test_short_text_with_na(self):
        sections = {"SECTION F": "N/A"}
        section_map = {"SECTION F": "SECTION F"}
        assert _check_not_applicable_required_when_blank(self._rule(), sections, section_map) is None

    def test_short_text_with_not_applicable(self):
        sections = {"SECTION F": "Not Applicable"}
        section_map = {"SECTION F": "SECTION F"}
        assert _check_not_applicable_required_when_blank(self._rule(), sections, section_map) is None

    def test_short_text_no_na_raises_finding(self):
        sections = {"SECTION F": "TBD"}
        section_map = {"SECTION F": "SECTION F"}
        finding = _check_not_applicable_required_when_blank(self._rule(), sections, section_map)
        assert finding is not None
        assert "fewer than 40 characters" in finding.message
        assert finding.severity == "WARNING"

    def test_empty_text_no_na_raises_finding(self):
        sections = {"SECTION F": ""}
        section_map = {"SECTION F": "SECTION F"}
        finding = _check_not_applicable_required_when_blank(self._rule(), sections, section_map)
        assert finding is not None

    def test_section_missing_no_finding(self):
        sections = {}
        section_map = {"SECTION F": None}
        assert _check_not_applicable_required_when_blank(self._rule(), sections, section_map) is None


# ---------------------------------------------------------------------------
# _check_must_mention_keywords
# ---------------------------------------------------------------------------

class TestMustMentionKeywords:
    def _rule(self, keywords=None, min_hits=1, section="SECTION D"):
        return {
            "id": "K001", "type": "must_mention_keywords",
            "section": section, "severity": "ERROR",
            "keywords": keywords or ["sampling", "random", "sample size"],
            "min_hits": min_hits,
        }

    def test_enough_hits(self):
        sections = {"SECTION D": "We used sampling with random selection."}
        section_map = {"SECTION D": "SECTION D"}
        assert _check_must_mention_keywords(self._rule(min_hits=2), sections, section_map) is None

    def test_exact_min_hits(self):
        sections = {"SECTION D": "Sampling was applied."}
        section_map = {"SECTION D": "SECTION D"}
        assert _check_must_mention_keywords(self._rule(min_hits=1), sections, section_map) is None

    def test_not_enough_hits(self):
        sections = {"SECTION D": "Data was collected."}
        section_map = {"SECTION D": "SECTION D"}
        finding = _check_must_mention_keywords(self._rule(min_hits=2), sections, section_map)
        assert finding is not None
        assert "0/2" in finding.message
        assert finding.severity == "ERROR"

    def test_missing_keywords_listed(self):
        sections = {"SECTION D": "We used sampling."}
        section_map = {"SECTION D": "SECTION D"}
        finding = _check_must_mention_keywords(self._rule(min_hits=2), sections, section_map)
        assert finding is not None
        assert "random" in finding.message
        assert "sample size" in finding.message

    def test_case_insensitive(self):
        sections = {"SECTION D": "SAMPLING and RANDOM selection."}
        section_map = {"SECTION D": "SECTION D"}
        assert _check_must_mention_keywords(self._rule(min_hits=2), sections, section_map) is None

    def test_section_missing_no_finding(self):
        sections = {}
        section_map = {"SECTION D": None}
        assert _check_must_mention_keywords(self._rule(), sections, section_map) is None

    def test_custom_message(self):
        rule = self._rule()
        rule["message"] = "Custom message about missing keywords"
        sections = {"SECTION D": "No relevant content."}
        section_map = {"SECTION D": "SECTION D"}
        finding = _check_must_mention_keywords(rule, sections, section_map)
        assert finding is not None
        assert finding.message == "Custom message about missing keywords"

    def test_aggregated_text_across_duplicates(self):
        sections = {
            "section d. data and parameters": "We collected sampling data.",
            "SECTION D. DATA AND PARAMETERS": "Using random selection method.",
        }
        section_map = {"SECTION D": "SECTION D. DATA AND PARAMETERS"}
        finding = _check_must_mention_keywords(self._rule(min_hits=2), sections, section_map)
        assert finding is None

    def test_text_normalization_with_newlines(self):
        sections = {"SECTION D": "We used\nsampling\n\nwith random\r\nselection."}
        section_map = {"SECTION D": "SECTION D"}
        assert _check_must_mention_keywords(self._rule(min_hits=2), sections, section_map) is None

    def test_confidence_precision_keywords(self):
        rule = self._rule(keywords=["90/10", "confidence", "precision"], min_hits=1)
        sections = {"SECTION D": "Confidence level of 95%."}
        section_map = {"SECTION D": "SECTION D"}
        assert _check_must_mention_keywords(rule, sections, section_map) is None

    def test_kpt_keyword_warning(self):
        rule = {
            "id": "K003", "type": "must_mention_keywords",
            "section": "SECTION D", "severity": "WARNING",
            "keywords": ["KPT", "Kitchen Performance Test"], "min_hits": 1,
        }
        sections = {"SECTION D": "KPT results show improvement."}
        section_map = {"SECTION D": "SECTION D"}
        assert _check_must_mention_keywords(rule, sections, section_map) is None

    def test_grievance_keywords(self):
        rule = self._rule(keywords=["grievance", "complaint", "stakeholder"], section="SECTION F")
        sections = {"SECTION F": "Stakeholder consultation was conducted."}
        section_map = {"SECTION F": "SECTION F"}
        assert _check_must_mention_keywords(rule, sections, section_map) is None


# ---------------------------------------------------------------------------
# compliance_score
# ---------------------------------------------------------------------------

class TestComplianceScore:
    def test_no_findings(self):
        assert compute_compliance_score([]) == 100

    def test_one_error(self):
        assert compute_compliance_score([Finding(rule_id="X", severity="ERROR", category="STRUCTURE", message="m")]) == 90

    def test_one_warning(self):
        assert compute_compliance_score([Finding(rule_id="X", severity="WARNING", category="FORMAT", message="m")]) == 97

    def test_mixed(self):
        findings = [
            Finding(rule_id="A", severity="ERROR", category="STRUCTURE", message="m"),
            Finding(rule_id="B", severity="ERROR", category="KEY_FIELDS", message="m"),
            Finding(rule_id="C", severity="WARNING", category="FORMAT", message="m"),
        ]
        assert compute_compliance_score(findings) == 77

    def test_floor_at_zero(self):
        findings = [Finding(rule_id=f"E{i}", severity="ERROR", category="STRUCTURE", message="m") for i in range(15)]
        assert compute_compliance_score(findings) == 0

    def test_info_no_penalty(self):
        assert compute_compliance_score([Finding(rule_id="I", severity="INFO", category="FORMAT", message="m")]) == 100


# ---------------------------------------------------------------------------
# compute_status
# ---------------------------------------------------------------------------

class TestComputeStatus:
    def test_pass_no_findings(self):
        assert compute_status([]) == "PASS"

    def test_fail_structure_error(self):
        findings = [Finding(rule_id="S1", severity="ERROR", category="STRUCTURE", message="m")]
        assert compute_status(findings) == "FAIL"

    def test_fail_key_fields_error(self):
        findings = [Finding(rule_id="KF1", severity="ERROR", category="KEY_FIELDS", message="m")]
        assert compute_status(findings) == "FAIL"

    def test_review_structure_warning_only(self):
        findings = [Finding(rule_id="S1", severity="WARNING", category="STRUCTURE", message="m")]
        assert compute_status(findings) == "REVIEW"

    def test_review_format_error(self):
        findings = [Finding(rule_id="F1", severity="ERROR", category="FORMAT", message="m")]
        assert compute_status(findings) == "REVIEW"

    def test_review_content_hint_error(self):
        findings = [Finding(rule_id="C1", severity="ERROR", category="CONTENT_HINT", message="m")]
        assert compute_status(findings) == "REVIEW"

    def test_review_info_only(self):
        findings = [Finding(rule_id="I1", severity="INFO", category="FORMAT", message="m")]
        assert compute_status(findings) == "REVIEW"

    def test_fail_mixed_with_structure_error(self):
        findings = [
            Finding(rule_id="S1", severity="ERROR", category="STRUCTURE", message="m"),
            Finding(rule_id="F1", severity="WARNING", category="FORMAT", message="m"),
            Finding(rule_id="C1", severity="ERROR", category="CONTENT_HINT", message="m"),
        ]
        assert compute_status(findings) == "FAIL"

    def test_review_mixed_without_critical_error(self):
        findings = [
            Finding(rule_id="F1", severity="ERROR", category="FORMAT", message="m"),
            Finding(rule_id="C1", severity="WARNING", category="CONTENT_HINT", message="m"),
        ]
        assert compute_status(findings) == "REVIEW"

    def test_review_key_fields_warning(self):
        findings = [Finding(rule_id="KF1", severity="WARNING", category="KEY_FIELDS", message="m")]
        assert compute_status(findings) == "REVIEW"


# ---------------------------------------------------------------------------
# End-to-end with real docx
# ---------------------------------------------------------------------------

class TestEndToEnd:
    @staticmethod
    def _make_docx(sections: dict[str, str]) -> str:
        doc = Document()
        for heading, body in sections.items():
            doc.add_heading(heading, level=1)
            doc.add_paragraph(body)
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        return tmp.name

    def test_full_pipeline_with_fields_and_dates(self):
        path = self._make_docx({
            "KEY PROJECT INFORMATION": (
                "GS ID: GS-1234 Title of the project: Solar Cookstoves Kenya "
                "Version number of the monitoring report: v2.0 "
                "Completion date: 15/06/2024 "
                "Monitoring period number: MP #3 "
                "Duration of this monitoring period: from 01/01/2024 to 31/12/2024 "
                "Host Country: Kenya "
                "Methodology applied: AMS-II.G version 09"
            ),
            "SECTION A. DESCRIPTION OF PROJECT": "Solar cookstove distribution in rural Kenya.",
            "SECTION B. IMPLEMENTATION OF PROJECT": "Implemented across 5 counties.",
            "SECTION C. MONITORING SYSTEM": "Continuous monitoring with data source from spreadsheet sensors.",
            "SECTION D. DATA AND PARAMETERS": (
                "Parameters measured quarterly using sampling and random selection. "
                "Confidence level 90/10. KPT Kitchen Performance Test conducted."
            ),
            "SECTION E. CALCULATION OF SDG IMPACTS": (
                "Baseline emissions: 15,000 tCO2e. "
                "Project emissions: 500 tCO2e. "
                "Emission reductions: 14,500 tCO2e net reductions."
            ),
            "SECTION F. SAFEGUARDS REPORTING": (
                "Environmental and social safeguards assessment completed. "
                "Grievance mechanism in place. Stakeholder consultation conducted."
            ),
            "SECTION G. STAKEHOLDER INPUTS": "No disputes reported during the monitoring period for the project.",
        })

        result = run_analysis(path)
        assert result.compliant is True
        assert result.compliance_score == 100
        assert len(result.findings) == 0
        assert result.status == "PASS"
        assert result.status_label == "BASIC CHECKS PASSED"

    def test_missing_fields_and_bad_dates(self):
        path = self._make_docx({
            "KEY PROJECT INFORMATION": (
                "Some basic info. Completion date: 2024-06-15."
            ),
            "SECTION A. DESCRIPTION": "A project.",
            "SECTION B. IMPLEMENTATION": "Implemented.",
            "SECTION C. MONITORING": "Monitoring.",
            "SECTION D. DATA": "Data.",
            "SECTION E. SDG IMPACTS": "Baseline emissions TBD. Project emissions TBD.",
            "SECTION F. SAFEGUARDS": "TBD",
            "SECTION G. STAKEHOLDER": "N/A",
        })

        result = run_analysis(path)
        assert result.compliant is False
        assert result.compliance_score < 100
        assert result.status == "FAIL"
        assert result.status_label == "NOT READY FOR SUBMISSION"

        messages = [f.message for f in result.findings]
        assert any("gs_id" in m for m in messages)
        assert any("title_of_project" in m for m in messages)
        assert any("Date format violation" in m for m in messages)
        assert any("fewer than 40 characters" in m for m in messages)

        categories = [f.category for f in result.findings]
        assert "KEY_FIELDS" in categories
        assert "FORMAT" in categories
