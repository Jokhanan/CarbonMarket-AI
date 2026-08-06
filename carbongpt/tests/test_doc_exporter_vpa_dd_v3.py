"""
Tests for carbongpt.core.doc_exporter's v1.0 VPA-DD v3.0 export path
(_fill_gs_vpa_dd_v3 and helpers) — the export side of "génère un VPA-DD
complet... exporte-le en Word sur le template v3.0". Runs against the real
v3.0 .docx file already tracked in document_repository/ — no network, no
database except for the one DB-gated resolver test.
"""
import os
from pathlib import Path

import docx
import pytest

from carbongpt.core.doc_exporter import (
    _fill_gs_vpa_dd_v3,
    _fill_parameter_block_field,
    _parse_field_value_block,
    _resolve_gs_vpa_dd_v3_template,
)

REPO_DIR = Path(__file__).parent.parent.parent / "document_repository"
V3_0_PATH = REPO_DIR / "85924c132bed502d_T-PAA_PreReview_V3.0_VPA-Design-Document.docx"

requires_v3_0 = pytest.mark.skipif(not V3_0_PATH.exists(), reason="VPA-DD v3.0 not ingested yet")


class TestParseFieldValueBlock:
    def test_splits_field_colon_value_lines(self):
        text = "Data/parameter: WCCF\nUnit: Ratio (kg wood per kg charcoal)\nDescription: A factor."
        fields = _parse_field_value_block(text)
        assert fields["data/parameter"] == "WCCF"
        assert fields["unit"] == "Ratio (kg wood per kg charcoal)"

    def test_ignores_lines_without_a_colon(self):
        text = "Data/parameter: WCCF\nThis is a plain sentence with no label."
        fields = _parse_field_value_block(text)
        assert len(fields) == 1


@requires_v3_0
class TestFillParameterBlockField:
    def test_fills_original_table_and_inserts_one_copy_per_extra_instance(self):
        doc = docx.Document(str(V3_0_PATH))
        before_table_count = len(doc.tables)
        combined = (
            "=== PARAMETER BLOCK: ICS 17 ===\n"
            "Data/parameter: WCCF\nUnit: Ratio (kg wood per kg charcoal)\n\n"
            "=== PARAMETER BLOCK: ICS 18 ===\n"
            "Data/parameter: fNRB\nUnit: Fraction"
        )
        _fill_parameter_block_field(doc.tables[37], combined)

        assert len(doc.tables) == before_table_count + 1  # one extra instance inserted
        first_table = doc.tables[37]
        second_table = doc.tables[38]
        first_row_texts = [r.cells[1].text for r in first_table.rows if "data/parameter" in r.cells[0].text.lower()]
        second_row_texts = [r.cells[1].text for r in second_table.rows if "data/parameter" in r.cells[0].text.lower()]
        assert first_row_texts == ["WCCF"]
        assert second_row_texts == ["fNRB"]

    def test_not_available_marker_leaves_table_unfilled(self):
        doc = docx.Document(str(V3_0_PATH))
        table = doc.tables[55]
        before = table.rows[0].cells[1].text  # official template placeholder, untouched
        _fill_parameter_block_field(table, "=== PARAMETER BLOCK: n/a ===\n[NOT AVAILABLE: no source]")
        after = table.rows[0].cells[1].text
        assert after == before


@requires_v3_0
class TestFillGsVpaDdV3:
    def test_inserts_prose_content_after_the_matching_heading(self):
        doc = docx.Document(str(V3_0_PATH))
        heading_para = doc.paragraphs[96]
        assert heading_para.text.strip() == "Application of methodology (ies)"

        _fill_gs_vpa_dd_v3(doc, {"H96": "This is the drafted introduction for this section."}, {"name": "Test"})

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "This is the drafted introduction for this section." in full_text

    def test_two_parameter_block_sections_do_not_shift_each_other(self):
        # Regression (found generating a real document, v1.0, 04.08.2026):
        # T37 has 10 rows in this file, T40 has 18 — real, distinguishable
        # shapes. Filling T37 first with several extra instances used to
        # shift every later table's position, so T40's own content landed
        # in whatever unrelated table now sat at the STALE index 40. Both
        # must end up correctly filled regardless of processing order.
        doc = docx.Document(str(V3_0_PATH))
        assert len(doc.tables[37].rows) == 10
        assert len(doc.tables[40].rows) == 18

        _fill_gs_vpa_dd_v3(
            doc,
            {
                "T37": (
                    "=== PARAMETER BLOCK: ICS 17 ===\nData/parameter: WCCF\n\n"
                    "=== PARAMETER BLOCK: ICS 18 ===\nData/parameter: fNRB\n\n"
                    "=== PARAMETER BLOCK: ICS 1 ===\nData/parameter: EF_kiln"
                ),
                "T40": "=== PARAMETER BLOCK: ICS 22 ===\nData/parameter: Usage rate",
            },
            {"name": "Test"},
        )

        full_text = "\n".join(
            "\t".join(c.text for c in row.cells) for table in doc.tables for row in table.rows
        )
        assert "WCCF" in full_text
        assert "fNRB" in full_text
        assert "EF_kiln" in full_text
        assert "Usage rate" in full_text
        # The 18-row monitoring table must still exist, correctly filled —
        # not silently overwritten or duplicated into a 10-row table.
        eighteen_row_tables = [t for t in doc.tables if len(t.rows) == 18]
        assert any(
            any("Usage rate" in c.text for row in t.rows for c in row.cells)
            for t in eighteen_row_tables
        )

    def test_fills_parameter_block_table_for_a_t_prefixed_key(self):
        doc = docx.Document(str(V3_0_PATH))
        _fill_gs_vpa_dd_v3(
            doc,
            {"T37": "=== PARAMETER BLOCK: ICS 17 ===\nData/parameter: WCCF"},
            {"name": "Test"},
        )
        table = doc.tables[37]
        row_texts = [r.cells[1].text for r in table.rows if "data/parameter" in r.cells[0].text.lower()]
        assert row_texts == ["WCCF"]

    def test_ignores_out_of_range_or_malformed_field_keys(self):
        doc = docx.Document(str(V3_0_PATH))
        # Must not raise for a field_key with no counterpart in this file.
        _fill_gs_vpa_dd_v3(doc, {"H99999": "orphan content", "Xbad": "bad prefix"}, {"name": "Test"})

    def test_empty_content_is_skipped(self):
        doc = docx.Document(str(V3_0_PATH))
        _fill_gs_vpa_dd_v3(doc, {"H96": "   "}, {"name": "Test"})
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # No new non-heading paragraph was inserted right after H96 with blank content.
        assert full_text.count("Application of methodology (ies)") == 1


def _db_available() -> bool:
    if not os.getenv("DATABASE_URL"):
        return False
    try:
        from carbongpt.repository.db import get_cursor
        with get_cursor() as cur:
            cur.execute("SELECT 1")
        return True
    except Exception:
        return False


requires_db = pytest.mark.skipif(not _db_available(), reason="DATABASE_URL not set or DB unreachable")


@requires_db
class TestResolveGsVpaDdV3Template:
    def test_returns_a_real_existing_file_when_analyzed(self):
        local_path, template_version_id = _resolve_gs_vpa_dd_v3_template()
        if local_path is None:
            pytest.skip("VPA-DD v3.0 not analyzed in this database")
        assert os.path.isfile(local_path)
        assert isinstance(template_version_id, int)

    def test_never_raises_on_db_outage(self, monkeypatch):
        import carbongpt.repository.db as db_module

        def _raise(*args, **kwargs):
            raise RuntimeError("simulated database outage")

        monkeypatch.setattr(db_module, "get_cursor", _raise)
        assert _resolve_gs_vpa_dd_v3_template() == (None, None)
