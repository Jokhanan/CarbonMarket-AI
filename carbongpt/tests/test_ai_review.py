"""
test_ai_review.py — Tests for AI review guide structure and response validation.
"""

import json
import tempfile
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from carbongpt.guides.gs_mr_perfcert_v1_2 import (
    GUIDE_ID,
    get_subsections,
    get_subsection,
    get_parent_sections,
    get_subsections_for_parent,
)
from carbongpt.core.models import (
    AIReviewRequest,
    AIReviewResponse,
    SectionReview,
    GlobalSummary,
)
from carbongpt.core.ai_review import (
    _build_section_prompt,
    _build_global_prompt,
    review_section,
    review_global,
    run_ai_review,
    SECTION_REVIEW_SCHEMA,
    GLOBAL_SUMMARY_SCHEMA,
)


class TestGuideStructure:
    def test_guide_id(self):
        assert GUIDE_ID == "GoldStandard_MR_PerfCert_v1_2"

    def test_subsections_not_empty(self):
        subs = get_subsections()
        assert len(subs) >= 19

    def test_all_subsections_have_required_keys(self):
        required_keys = {"title", "parent_section", "must_include", "examples", "failure_modes"}
        for sub_id, sub in get_subsections().items():
            missing = required_keys - set(sub.keys())
            assert not missing, f"Subsection {sub_id} missing keys: {missing}"

    def test_all_must_include_non_empty(self):
        for sub_id, sub in get_subsections().items():
            assert len(sub["must_include"]) > 0, f"{sub_id} has empty must_include"

    def test_all_failure_modes_non_empty(self):
        for sub_id, sub in get_subsections().items():
            assert len(sub["failure_modes"]) > 0, f"{sub_id} has empty failure_modes"

    def test_section_a_subsections(self):
        a_subs = get_subsections_for_parent("SECTION A")
        assert set(a_subs.keys()) == {"A.1", "A.2", "A.3", "A.4"}

    def test_section_b_subsections(self):
        b_subs = get_subsections_for_parent("SECTION B")
        assert set(b_subs.keys()) == {"B.1", "B.2", "B.3"}

    def test_get_subsection_exists(self):
        sub = get_subsection("A.1")
        assert sub is not None
        assert sub["title"] == "General description of the project"

    def test_get_subsection_not_found(self):
        assert get_subsection("Z.99") is None

    def test_parent_sections(self):
        parents = get_parent_sections()
        for s in ["SECTION A", "SECTION B", "SECTION C", "SECTION D", "SECTION E", "SECTION F", "SECTION G"]:
            assert s in parents

    def test_section_c_subsections(self):
        c_subs = get_subsections_for_parent("SECTION C")
        assert set(c_subs.keys()) == {"C.1"}

    def test_section_d_subsections(self):
        d_subs = get_subsections_for_parent("SECTION D")
        assert set(d_subs.keys()) == {"D.1", "D.2", "D.3", "D.4"}

    def test_section_e_subsections(self):
        e_subs = get_subsections_for_parent("SECTION E")
        assert set(e_subs.keys()) == {"E.1", "E.2", "E.3", "E.4", "E.5", "E.6"}

    def test_section_f_subsections(self):
        f_subs = get_subsections_for_parent("SECTION F")
        assert set(f_subs.keys()) == {"F.1"}

    def test_section_g_subsections(self):
        g_subs = get_subsections_for_parent("SECTION G")
        assert set(g_subs.keys()) == {"G.1", "G.2", "G.3"}


class TestResponseModels:
    def test_section_review_model(self):
        review = SectionReview(
            section_id="A.1",
            section_title="General description",
            completeness_score=75,
            issues=["Missing GHG mechanism"],
            suggested_fixes=["[DRAFT] Add description of emission reduction mechanism."],
            questions_for_user=["What type of technology is used?"],
        )
        assert review.section_id == "A.1"
        assert review.completeness_score == 75
        assert len(review.issues) == 1

    def test_global_summary_model(self):
        summary = GlobalSummary(
            overall_risk="MEDIUM",
            top_issues=["Missing coordinates"],
            top_actions=["Add GPS coordinates to A.2"],
            coherence_flags=["Section A and B dates do not align"],
        )
        assert summary.overall_risk == "MEDIUM"

    def test_ai_review_response_model(self):
        resp = AIReviewResponse(
            per_section_reviews=[
                SectionReview(
                    section_id="A.1",
                    section_title="General description",
                    completeness_score=80,
                ),
            ],
            global_summary=GlobalSummary(
                overall_risk="LOW",
            ),
        )
        assert len(resp.per_section_reviews) == 1
        assert resp.global_summary.overall_risk == "LOW"

    def test_ai_review_request_model(self):
        req = AIReviewRequest(
            standard="GoldStandard",
            doc_type="MR",
            version="PerfCert_v1_2",
            doc_path="/tmp/test.docx",
        )
        assert req.standard == "GoldStandard"

    def test_section_review_score_bounds(self):
        with pytest.raises(Exception):
            SectionReview(
                section_id="A.1",
                section_title="Test",
                completeness_score=150,
            )

    def test_global_summary_risk_validation(self):
        with pytest.raises(Exception):
            GlobalSummary(overall_risk="CRITICAL")


class TestPromptBuilding:
    def test_section_prompt_contains_requirements(self):
        guide = get_subsection("A.1")
        prompt = _build_section_prompt("A.1", guide, "Some project text here.")
        assert "A.1" in prompt
        assert "General description" in prompt
        assert "must include" in prompt.lower()
        assert "Some project text here." in prompt

    def test_section_prompt_contains_failure_modes(self):
        guide = get_subsection("A.2")
        prompt = _build_section_prompt("A.2", guide, "Location info.")
        assert "Failure Modes" in prompt
        assert "no geographic coordinates" in prompt.lower()

    def test_global_prompt_contains_section_results(self):
        reviews = [
            {"section_id": "A.1", "section_title": "General description", "completeness_score": 80, "issues": ["Missing mechanism"]},
            {"section_id": "A.2", "section_title": "Location", "completeness_score": 60, "issues": ["No coordinates"]},
        ]
        prompt = _build_global_prompt(reviews)
        assert "A.1" in prompt
        assert "A.2" in prompt
        assert "score=80" in prompt
        assert "Missing mechanism" in prompt


class TestSchemas:
    def test_section_review_schema_has_required_fields(self):
        assert "completeness_score" in SECTION_REVIEW_SCHEMA["properties"]
        assert "issues" in SECTION_REVIEW_SCHEMA["properties"]
        assert "suggested_fixes" in SECTION_REVIEW_SCHEMA["properties"]
        assert "questions_for_user" in SECTION_REVIEW_SCHEMA["properties"]

    def test_global_summary_schema_has_required_fields(self):
        assert "overall_risk" in GLOBAL_SUMMARY_SCHEMA["properties"]
        assert "top_issues" in GLOBAL_SUMMARY_SCHEMA["properties"]
        assert "top_actions" in GLOBAL_SUMMARY_SCHEMA["properties"]
        assert "coherence_flags" in GLOBAL_SUMMARY_SCHEMA["properties"]


def _mock_openai_json(content_dict: dict):
    return {
        "choices": [
            {"message": {"content": json.dumps(content_dict)}}
        ]
    }


class TestReviewWithMockedLLM:
    def test_review_section_mocked(self):
        mock_response = _mock_openai_json({
            "completeness_score": 65,
            "issues": ["No GHG mechanism described"],
            "suggested_fixes": ["[DRAFT] Add description of how the project reduces emissions."],
            "questions_for_user": ["What technology is used for emission reduction?"],
        })

        with patch("carbongpt.core.ai_review._call_openai_structured", return_value=mock_response["choices"][0]["message"]["content"]):
            pass

        with patch("carbongpt.core.ai_review._call_openai_structured") as mock_call:
            mock_call.return_value = {
                "completeness_score": 65,
                "issues": ["No GHG mechanism described"],
                "suggested_fixes": ["[DRAFT] Add description of how the project reduces emissions."],
                "questions_for_user": ["What technology is used for emission reduction?"],
            }
            guide = get_subsection("A.1")
            result = review_section("fake-key", "A.1", guide, "A cookstove project.")

            assert result["section_id"] == "A.1"
            assert result["completeness_score"] == 65
            assert len(result["issues"]) == 1
            assert "[DRAFT]" in result["suggested_fixes"][0]
            mock_call.assert_called_once()

    def test_review_global_mocked(self):
        with patch("carbongpt.core.ai_review._call_openai_structured") as mock_call:
            mock_call.return_value = {
                "overall_risk": "MEDIUM",
                "top_issues": ["Missing coordinates in A.2"],
                "top_actions": ["Add GPS coordinates"],
                "coherence_flags": ["Project start date in B.1 not consistent with A.1"],
            }

            reviews = [
                {"section_id": "A.1", "section_title": "General description", "completeness_score": 80, "issues": []},
                {"section_id": "A.2", "section_title": "Location", "completeness_score": 40, "issues": ["No coordinates"]},
            ]
            result = review_global("fake-key", reviews)

            assert result["overall_risk"] == "MEDIUM"
            assert len(result["top_issues"]) == 1
            assert len(result["coherence_flags"]) == 1
            mock_call.assert_called_once()

    def test_run_ai_review_full_mocked(self):
        doc = Document()
        doc.add_heading("SECTION A. DESCRIPTION OF PROJECT", level=1)
        doc.add_paragraph(
            "This project distributes improved cookstoves in Siaya County, Kenya "
            "(0.0617 S, 34.2422 E). Project Developer: CleanCook Ltd. "
            "Methodology: AMS-II.G version 09."
        )
        doc.add_heading("SECTION B. IMPLEMENTATION OF PROJECT", level=1)
        doc.add_paragraph(
            "The project has been operational since 01/01/2020. No post-registration changes. "
            "Complies with all Gold Standard requirements."
        )
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)

        section_result = {
            "completeness_score": 70,
            "issues": ["Incomplete description"],
            "suggested_fixes": ["[DRAFT] Expand description"],
            "questions_for_user": ["What is the project capacity?"],
        }

        global_result = {
            "overall_risk": "LOW",
            "top_issues": ["Minor gaps"],
            "top_actions": ["Complete descriptions"],
            "coherence_flags": [],
        }

        total_subsections = len(get_subsections())
        reviewed_count = len(get_subsections_for_parent("SECTION A")) + len(get_subsections_for_parent("SECTION B"))
        call_count = [0]

        def side_effect(*args, **kwargs):
            call_count[0] += 1
            if call_count[0] <= reviewed_count:
                return section_result
            return global_result

        with patch("carbongpt.core.ai_review._call_openai_structured", side_effect=side_effect), \
             patch("carbongpt.core.ai_review._get_api_key", return_value="fake-key"):
            result = run_ai_review(tmp.name)

        assert "per_section_reviews" in result
        assert "global_summary" in result
        assert len(result["per_section_reviews"]) == total_subsections
        assert result["global_summary"]["overall_risk"] == "LOW"

        for review in result["per_section_reviews"]:
            assert "section_id" in review
            assert "completeness_score" in review

    def test_missing_sections_get_score_zero(self):
        doc = Document()
        doc.add_heading("SECTION B. IMPLEMENTATION", level=1)
        doc.add_paragraph("Operational since 2020.")
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)

        section_result = {
            "completeness_score": 60,
            "issues": [],
            "suggested_fixes": [],
            "questions_for_user": [],
        }
        global_result = {
            "overall_risk": "HIGH",
            "top_issues": ["Section A missing"],
            "top_actions": ["Add Section A"],
            "coherence_flags": [],
        }

        call_count = [0]
        b_subsection_count = len(get_subsections_for_parent("SECTION B"))

        def side_effect(*args, **kwargs):
            call_count[0] += 1
            if call_count[0] <= b_subsection_count:
                return section_result
            return global_result

        with patch("carbongpt.core.ai_review._call_openai_structured", side_effect=side_effect), \
             patch("carbongpt.core.ai_review._get_api_key", return_value="fake-key"):
            result = run_ai_review(tmp.name)

        a_reviews = [r for r in result["per_section_reviews"] if r["section_id"].startswith("A.")]
        for r in a_reviews:
            assert r["completeness_score"] == 0
            assert any("not found" in issue.lower() for issue in r["issues"])
