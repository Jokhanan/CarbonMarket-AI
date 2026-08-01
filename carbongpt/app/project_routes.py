import logging
import os
import uuid
from pathlib import Path

from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from carbongpt.app.config import BASE_DIR

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/projects", tags=["projects"])

PROJECT_FILES_DIR = BASE_DIR / "project_files"
PROJECT_FILES_DIR.mkdir(parents=True, exist_ok=True)


class ProjectCreate(BaseModel):
    name: str
    standard: str
    doc_type: str | None = None
    methodology: str | None = None
    country: str | None = None
    description: str | None = None
    project_type: str | None = None
    parent_project_id: int | None = None
    monitoring_period_start: str | None = None
    monitoring_period_end: str | None = None
    methodology_settings: dict | None = None
    location_name: str | None = None
    region: str | None = None
    district: str | None = None
    latitude: float | None = None
    longitude: float | None = None
    boundary_geojson: str | None = None
    project_intake: dict | None = None


class ProjectUpdate(BaseModel):
    name: str | None = None
    standard: str | None = None
    doc_type: str | None = None
    methodology: str | None = None
    country: str | None = None
    description: str | None = None
    status: str | None = None
    crediting_period_start: str | None = None
    crediting_period_years: int | None = None
    project_settings: dict | None = None
    project_intake: dict | None = None
    methodology_settings: dict | None = None
    location_name: str | None = None
    region: str | None = None
    district: str | None = None
    latitude: float | None = None
    longitude: float | None = None
    boundary_geojson: str | None = None
    project_type: str | None = None
    parent_project_id: int | None = None
    monitoring_period_start: str | None = None
    monitoring_period_end: str | None = None


class MonitoringPeriodCreate(BaseModel):
    period_number: int = 1
    period_start: str | None = None
    period_end: str | None = None
    status: str = "planned"
    notes: str | None = None


class MonitoringPeriodUpdate(BaseModel):
    period_number: int | None = None
    period_start: str | None = None
    period_end: str | None = None
    status: str | None = None
    notes: str | None = None
    mr_project_id: int | None = None


class WriteSectionRequest(BaseModel):
    section_id: str
    user_instructions: str | None = None


class WriteAllRequest(BaseModel):
    user_instructions: str | None = None


class UpdateSectionTextRequest(BaseModel):
    section_id: str
    doc_type: str = "pdd"
    text: str


class ExplainSectionRequest(BaseModel):
    section_id: str


class ParseMethodologyRequest(BaseModel):
    methodology_code: str


class RunCalculationRequest(BaseModel):
    method_id: str | None = None
    crediting_years: int = 7
    user_inputs: dict = Field(default_factory=dict)


class ExportCalculationRequest(BaseModel):
    calculation_result: dict = {}
    doc_type: str = "pdd"
    format: str = "excel"


class GenerateTemplateRequest(BaseModel):
    doc_type: str = "pdd"
    sections_to_generate: list[str] | None = None
    include_calculations: bool = True
    calculation_result: dict | None = None
    user_instructions: str | None = None


class ReviewDocumentRequest(BaseModel):
    document_id: int


@router.get("/methodologies")
def get_methodologies(standard: str = None, category: str = None, search: str = None, limit: int = 200):
    from carbongpt.repository.store import list_methodologies
    return list_methodologies(standard=standard, category=category, search=search, limit=limit)


@router.get("/methodologies/categories")
def get_methodology_categories_endpoint():
    from carbongpt.repository.store import get_methodology_categories
    return get_methodology_categories()


@router.get("/methodologies/{code}")
def get_methodology_detail(code: str):
    from carbongpt.repository.store import get_methodology
    m = get_methodology(code)
    if not m:
        raise HTTPException(status_code=404, detail="Methodology not found.")
    return m


@router.post("/methodologies/populate")
def populate_methodologies():
    from carbongpt.repository.methodology_db import populate_methodologies_from_projects
    count = populate_methodologies_from_projects()
    return {"populated": count}


@router.get("")
def list_projects(status: str = None):
    from carbongpt.repository.store import list_user_projects
    return list_user_projects(status=status)


@router.post("")
def create_project(data: ProjectCreate):
    from carbongpt.repository.store import create_user_project
    project_id = create_user_project(
        name=data.name,
        standard=data.standard,
        doc_type=data.doc_type,
        methodology=data.methodology,
        country=data.country,
        description=data.description,
        project_type=data.project_type,
        parent_project_id=data.parent_project_id,
        monitoring_period_start=data.monitoring_period_start,
        monitoring_period_end=data.monitoring_period_end,
        methodology_settings=data.methodology_settings,
        location_name=data.location_name,
        region=data.region,
        district=data.district,
        latitude=data.latitude,
        longitude=data.longitude,
        boundary_geojson=data.boundary_geojson,
        project_intake=data.project_intake,
    )
    project_dir = PROJECT_FILES_DIR / str(project_id)
    project_dir.mkdir(parents=True, exist_ok=True)
    return {"id": project_id, "message": "Project created."}


@router.get("/{project_id}")
def get_project(project_id: int):
    from carbongpt.repository.store import get_user_project
    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")
    return project


@router.patch("/{project_id}")
def update_project(project_id: int, data: ProjectUpdate):
    from carbongpt.repository.store import get_user_project, update_user_project
    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")
    update_user_project(project_id, **data.dict(exclude_none=True))
    return {"message": "Project updated."}


@router.delete("/{project_id}")
def delete_project(project_id: int):
    from carbongpt.repository.store import delete_user_project
    delete_user_project(project_id)
    return {"message": "Project deleted."}


@router.get("/{project_id}/children")
def get_child_projects_endpoint(project_id: int):
    from carbongpt.repository.store import get_child_projects
    return get_child_projects(project_id)


@router.get("/{project_id}/monitoring-periods")
def list_monitoring_periods_endpoint(project_id: int):
    from carbongpt.repository.store import list_monitoring_periods
    return list_monitoring_periods(project_id)


@router.post("/{project_id}/monitoring-periods")
def create_monitoring_period_endpoint(project_id: int, data: MonitoringPeriodCreate):
    from carbongpt.repository.store import create_monitoring_period
    period_id = create_monitoring_period(
        project_id=project_id,
        period_number=data.period_number,
        period_start=data.period_start,
        period_end=data.period_end,
        status=data.status,
        notes=data.notes,
    )
    return {"id": period_id, "message": "Monitoring period created."}


@router.patch("/{project_id}/monitoring-periods/{period_id}")
def update_monitoring_period_endpoint(project_id: int, period_id: int, data: MonitoringPeriodUpdate):
    from carbongpt.repository.store import update_monitoring_period
    update_monitoring_period(period_id, **{k: v for k, v in data.dict().items() if v is not None})
    return {"message": "Monitoring period updated."}


@router.delete("/{project_id}/monitoring-periods/{period_id}")
def delete_monitoring_period_endpoint(project_id: int, period_id: int):
    from carbongpt.repository.store import delete_monitoring_period
    delete_monitoring_period(period_id)
    return {"message": "Monitoring period deleted."}


@router.post("/import-document")
async def import_document_endpoint(file: UploadFile = File(...)):
    import tempfile, json as _json
    suffix = Path(file.filename).suffix.lower() if file.filename else ".pdf"
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name

    parsed_text = ""
    try:
        if suffix in (".pdf",):
            try:
                import pdfplumber
                parts = []
                with pdfplumber.open(tmp_path) as pdf:
                    for page in pdf.pages[:30]:
                        t = page.extract_text()
                        if t:
                            parts.append(t)
                parsed_text = "\n".join(parts)
            except Exception as e:
                logger.warning("pdfplumber failed: %s", e)
        elif suffix in (".docx",):
            try:
                from docx import Document as DocxDoc
                doc_obj = DocxDoc(tmp_path)
                parsed_text = "\n".join(p.text.strip() for p in doc_obj.paragraphs if p.text.strip())
            except Exception as e:
                logger.warning("docx parse failed: %s", e)
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

    if not parsed_text:
        return {"error": "Could not extract text from this file.", "extracted": {}}

    excerpt = parsed_text[:8000]
    prompt = (
        "You are a carbon project expert. The following is text extracted from a carbon project document "
        "(e.g. PDD, PoA-DD, Monitoring Report).\n\n"
        "Extract the following fields and return ONLY valid JSON, no markdown:\n"
        "{\n"
        '  "project_name": "...",\n'
        '  "standard": "GoldStandard" or "Verra" or null,\n'
        '  "methodology": "exact methodology code e.g. TPDDTEC v4.0 or VM0050 or GS-MECD v1.2 or null",\n'
        '  "country": "country name or null",\n'
        '  "project_type": "standalone_pdd" or "poa_programme" or "monitoring_report" or null,\n'
        '  "activity_type": "Cooking devices" or "Renewable electricity" or "Other" or null,\n'
        '  "description": "one sentence summary or null",\n'
        '  "monitoring_period_start": "YYYY-MM-DD or null",\n'
        '  "monitoring_period_end": "YYYY-MM-DD or null"\n'
        "}\n\n"
        f"Document text:\n{excerpt}"
    )

    extracted = {}
    try:
        from carbongpt.core.ai_writer import _call_openai
        system_prompt = "You are a carbon project expert. Return ONLY valid JSON, no markdown fences."
        raw = _call_openai(system_prompt, prompt, max_tokens=600)
        if raw:
            raw = raw.strip()
            if raw.startswith("```"):
                raw = "\n".join(raw.split("\n")[1:])
            if raw.endswith("```"):
                raw = raw.rsplit("```", 1)[0]
            extracted = _json.loads(raw.strip())
    except Exception as e:
        logger.warning("AI extraction failed: %s", e)

    return {"extracted": extracted, "text_length": len(parsed_text)}


@router.patch("/{project_id}/documents/{doc_id}/ai-context")
def toggle_document_ai_context(project_id: int, doc_id: int, use_as_ai_context: bool = True):
    from carbongpt.repository.store import get_project_document, update_document_ai_context
    doc = get_project_document(doc_id)
    if not doc or doc["project_id"] != project_id:
        raise HTTPException(status_code=404, detail="Document not found.")
    update_document_ai_context(doc_id, use_as_ai_context)
    return {"message": "AI context updated.", "use_as_ai_context": use_as_ai_context}


@router.post("/{project_id}/documents")
def upload_project_document(
    project_id: int,
    file: UploadFile = File(...),
    doc_type: str = Form("other"),
    notes: str = Form(None),
):
    from carbongpt.repository.store import get_user_project, add_project_document, update_project_document

    VALID_DOC_TYPES = {"pdd", "mr", "valver", "poa_dd", "vpa_dd", "reference", "research", "field_data", "template", "other"}
    if doc_type not in VALID_DOC_TYPES:
        raise HTTPException(status_code=422, detail=f"Invalid doc_type '{doc_type}'. Must be one of: {', '.join(sorted(VALID_DOC_TYPES))}")

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else "other"
    file_type_map = {"pdf": "pdf", "docx": "docx", "xlsx": "xlsx", "csv": "csv"}
    file_type = file_type_map.get(ext, "other")

    project_dir = PROJECT_FILES_DIR / str(project_id)
    project_dir.mkdir(parents=True, exist_ok=True)

    safe_name = f"{uuid.uuid4().hex[:8]}_{file.filename}"
    file_path = project_dir / safe_name
    content = file.file.read()
    file_path.write_bytes(content)

    doc_id = add_project_document(
        project_id=project_id,
        doc_type=doc_type,
        file_name=file.filename,
        file_path=str(file_path),
        file_type=file_type,
        file_size_bytes=len(content),
        notes=notes,
    )

    parsed_text = None
    parsed_sections = []
    try:
        if file_type == "docx":
            from carbongpt.tools.parse_docx import parse_docx
            result = parse_docx(str(file_path))
            sections_dict = result.get("sections", {})
            if sections_dict:
                parsed_text = "\n\n".join(
                    f"{heading}\n{text}" if heading != "PREAMBLE" else text
                    for heading, text in sections_dict.items()
                    if text.strip()
                )
                parsed_sections = [
                    {"heading": heading, "text": text[:500]}
                    for heading, text in sections_dict.items()
                    if text.strip()
                ]
            if not parsed_text:
                from docx import Document as DocxDoc
                doc_obj = DocxDoc(str(file_path))
                all_text = []
                for para in doc_obj.paragraphs:
                    if para.text.strip():
                        all_text.append(para.text.strip())
                for table in doc_obj.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            all_text.append(" | ".join(row_text))
                parsed_text = "\n".join(all_text)
        elif file_type == "pdf":
            try:
                import pdfplumber
                text_parts = []
                with pdfplumber.open(str(file_path)) as pdf:
                    for page in pdf.pages:
                        t = page.extract_text()
                        if t:
                            text_parts.append(t)
                parsed_text = "\n".join(text_parts)
            except Exception as e:
                logger.warning("PDF parsing failed: %s", e)
    except Exception as e:
        logger.warning("Document parsing failed: %s", e)

    ai_summary = None
    if parsed_text:
        import json
        update_project_document(
            doc_id,
            parsed_text=parsed_text,
            parsed_sections=json.dumps(parsed_sections),
            status="parsed",
        )

        try:
            from carbongpt.core.ai_writer import extract_document_intelligence
            ai_summary = extract_document_intelligence(parsed_text, file.filename, doc_type)
            if ai_summary:
                update_project_document(doc_id, ai_extracted_summary=ai_summary)
        except Exception as e:
            logger.warning("AI extraction failed for %s: %s", file.filename, e)

        structured_data = None
        try:
            from carbongpt.core.ai_writer import extract_structured_intelligence
            structured_data = extract_structured_intelligence(parsed_text, file.filename, doc_type)
            if structured_data:
                update_project_document(doc_id, ai_extracted_data=structured_data)
        except Exception as e:
            logger.warning("Structured extraction failed for %s: %s", file.filename, e)

        try:
            from carbongpt.core.project_doc_index import index_project_document
            index_project_document(project_id, doc_id, parsed_text, parsed_sections, file.filename)
        except Exception as e:
            logger.warning("Project doc indexing failed for %s: %s", file.filename, e)

    return {
        "id": doc_id,
        "file_name": file.filename,
        "doc_type": doc_type,
        "parsed": parsed_text is not None,
        "ai_extracted": ai_summary is not None,
        "structured_extracted": structured_data is not None,
        "message": "Document uploaded successfully.",
    }


@router.delete("/{project_id}/documents/{doc_id}")
def delete_document(project_id: int, doc_id: int):
    from carbongpt.repository.store import get_project_document, delete_project_document
    doc = get_project_document(doc_id)
    if not doc or doc["project_id"] != project_id:
        raise HTTPException(status_code=404, detail="Document not found.")
    try:
        os.remove(doc["file_path"])
    except OSError:
        pass
    try:
        from carbongpt.core.project_doc_index import delete_project_doc_chunks
        delete_project_doc_chunks(doc_id)
    except Exception:
        pass
    delete_project_document(doc_id)
    return {"message": "Document deleted."}


@router.post("/{project_id}/documents/{doc_id}/extract-intelligence")
def extract_document_intel(project_id: int, doc_id: int):
    from carbongpt.repository.store import get_project_document, update_project_document
    doc = get_project_document(doc_id)
    if not doc or doc["project_id"] != project_id:
        raise HTTPException(status_code=404, detail="Document not found.")
    if not doc.get("parsed_text"):
        raise HTTPException(status_code=400, detail="Document has not been parsed yet.")
    try:
        from carbongpt.core.ai_writer import extract_document_intelligence
        summary = extract_document_intelligence(doc["parsed_text"], doc["file_name"], doc["doc_type"])
        if summary:
            update_project_document(doc_id, ai_extracted_summary=summary)

        structured_data = None
        try:
            from carbongpt.core.ai_writer import extract_structured_intelligence
            structured_data = extract_structured_intelligence(doc["parsed_text"], doc["file_name"], doc["doc_type"])
            if structured_data:
                update_project_document(doc_id, ai_extracted_data=structured_data)
        except Exception as e2:
            logger.warning("Structured extraction failed for doc %s: %s", doc_id, e2)

        if summary:
            return {"message": "Intelligence extracted.", "summary": summary, "structured_count": len(structured_data) if structured_data else 0}
        return {"message": "Extraction returned no results."}
    except Exception as e:
        err_str = str(e)
        logger.error("Intelligence extraction failed for doc %s: %s", doc_id, e)
        if "429" in err_str or "rate" in err_str.lower():
            raise HTTPException(status_code=429, detail="OpenAI rate limit reached. Please wait a minute and try again.")
        raise HTTPException(status_code=500, detail="Intelligence extraction failed. Please try again.")


@router.get("/{project_id}/intelligence-suggestions")
def get_intelligence_suggestions(project_id: int):
    from carbongpt.repository.store import get_user_project, get_project_documents_for_ai

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    docs = get_project_documents_for_ai(project_id)
    intake = project.get("project_intake") or {}
    dismissed = intake.get("_dismissed_suggestions", [])

    all_items = []
    for doc in docs:
        raw = doc.get("ai_extracted_data")
        if not raw:
            continue
        import json as _json
        if isinstance(raw, str):
            try:
                items = _json.loads(raw)
            except (ValueError, TypeError):
                continue
        else:
            items = raw
        if not isinstance(items, list):
            continue
        for item in items:
            if not isinstance(item, dict):
                continue
            item.setdefault("source", doc.get("file_name", ""))
            all_items.append(item)

    grouped = {}
    for item in all_items:
        cat = item.get("category", "")
        fkey = item.get("field_key", "")
        if not cat or not fkey:
            continue
        dismiss_key = f"{cat}.{fkey}"
        if dismiss_key in dismissed:
            continue

        if cat not in grouped:
            grouped[cat] = {}
        if fkey not in grouped[cat]:
            grouped[cat][fkey] = {
                "field_key": fkey,
                "values": [],
                "current_value": "",
            }
            cat_data = intake.get(cat, {})
            if isinstance(cat_data, dict):
                grouped[cat][fkey]["current_value"] = str(cat_data.get(fkey, ""))

        grouped[cat][fkey]["values"].append({
            "value": item.get("value", ""),
            "confidence": item.get("confidence", "medium"),
            "source": item.get("source", ""),
        })

    CONF_RANK = {"high": 0, "medium": 1, "low": 2}
    for cat in grouped:
        for fkey in grouped[cat]:
            grouped[cat][fkey]["values"].sort(key=lambda v: CONF_RANK.get(v.get("confidence", "medium"), 1))

    from carbongpt.core.ai_writer import INTAKE_FIELD_SCHEMA
    CATEGORY_LABELS = {
        "proponent": "Project Developer / Proponent",
        "project_overview": "Project Facts",
        "technology": "Technology",
        "location": "Location",
        "baseline_additionality": "Baseline & Additionality",
        "emission_reductions": "Emission Reductions",
        "monitoring": "Monitoring",
        "stakeholders": "Stakeholders",
        "safeguards": "Safeguards",
        "prior_consideration": "Prior Consideration",
        "legal_compliance": "Legal & Compliance",
        "monitoring_period": "Monitoring Period",
        "emission_factors": "Emission Factors & Parameters",
        "test_results": "Test Results",
    }

    result = []
    for cat, fields in grouped.items():
        cat_label = CATEGORY_LABELS.get(cat, cat.replace("_", " ").title())
        field_list = []
        schema_fields = INTAKE_FIELD_SCHEMA.get(cat, {})
        for fkey, fdata in fields.items():
            field_label = schema_fields.get(fkey, fkey.replace("_", " ").title())
            field_list.append({
                "field_key": fkey,
                "label": field_label,
                "values": fdata["values"],
                "current_value": fdata["current_value"],
            })
        result.append({
            "category": cat,
            "category_label": cat_label,
            "fields": field_list,
            "count": len(field_list),
        })

    result.sort(key=lambda x: x["count"], reverse=True)
    return {"suggestions": result, "total_count": sum(c["count"] for c in result)}


class IntelligenceConfirmRequest(BaseModel):
    items: list[dict]


@router.post("/{project_id}/intelligence-confirm")
def confirm_intelligence(project_id: int, body: IntelligenceConfirmRequest):
    from carbongpt.repository.store import get_user_project, update_user_project

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    intake = dict(project.get("project_intake") or {})
    sources = intake.get("_intelligence_sources", {})

    confirmed_count = 0
    skipped_count = 0
    for item in body.items:
        cat = item.get("category", "")
        fkey = item.get("field_key", "")
        val = item.get("value", "")
        source = item.get("source", "")
        force = item.get("force", False)
        if not cat or not fkey:
            continue
        if cat not in intake:
            intake[cat] = {}
        if isinstance(intake[cat], dict):
            existing = str(intake[cat].get(fkey, "")).strip()
            if existing and not force:
                skipped_count += 1
                continue
            intake[cat][fkey] = val
            sources[f"{cat}.{fkey}"] = source
            confirmed_count += 1

    intake["_intelligence_sources"] = sources
    update_user_project(project_id, project_intake=intake)
    msg = f"{confirmed_count} field(s) updated in Project Setup."
    if skipped_count > 0:
        msg += f" {skipped_count} field(s) skipped (already have values)."
    return {"message": msg, "confirmed": confirmed_count, "skipped": skipped_count}


class IntelligenceDismissRequest(BaseModel):
    items: list[dict]


@router.post("/{project_id}/intelligence-dismiss")
def dismiss_intelligence(project_id: int, body: IntelligenceDismissRequest):
    from carbongpt.repository.store import get_user_project, update_user_project

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    intake = dict(project.get("project_intake") or {})
    dismissed = list(intake.get("_dismissed_suggestions", []))

    dismissed_count = 0
    for item in body.items:
        cat = item.get("category", "")
        fkey = item.get("field_key", "")
        if not cat or not fkey:
            continue
        key = f"{cat}.{fkey}"
        if key not in dismissed:
            dismissed.append(key)
            dismissed_count += 1

    intake["_dismissed_suggestions"] = dismissed
    update_user_project(project_id, project_intake=intake)
    return {"message": f"{dismissed_count} suggestion(s) dismissed.", "dismissed": dismissed_count}


@router.get("/{project_id}/sections")
def get_document_sections(project_id: int, doc_type: str = "pdd"):
    from carbongpt.repository.store import get_user_project
    from carbongpt.core.ai_writer import get_sections_for_doc_type

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    sections = get_sections_for_doc_type(project["standard"], doc_type)
    if sections is None:
        raise HTTPException(status_code=400, detail=f"No guide available for {project['standard']}/{doc_type}")
    return sections


def _gather_ai_context(project_id, project, doc_type):
    from carbongpt.repository.store import get_project_documents_for_ai, get_project_documents_by_type
    pdd_text = None
    if doc_type == "mr":
        pdd_docs = get_project_documents_by_type(project_id, "pdd")
        if pdd_docs:
            pdd_text = pdd_docs[0].get("parsed_text", "")
        if not pdd_text and project.get("parent_project_id"):
            parent_pdd_docs = get_project_documents_by_type(project["parent_project_id"], "pdd")
            if parent_pdd_docs:
                pdd_text = parent_pdd_docs[0].get("parsed_text", "")

    DOC_TYPE_LABELS = {
        "pdd": "PDD", "mr": "Monitoring Report", "poa_dd": "PoA-DD",
        "vpa_dd": "VPA-DD", "valver": "Validation/Verification Report",
        "reference": "Reference Document", "research": "Research",
        "field_data": "Field Data / Survey", "other": "Supporting Document",
    }

    MAX_RAW_FALLBACK = 12000
    MAX_TOTAL_REF = 50000

    MAX_SUMMARY = 8000

    def _doc_text(rd, prefix="Document"):
        label = DOC_TYPE_LABELS.get(rd.get("doc_type", "other"), rd.get("doc_type", "other"))
        fname = rd.get("file_name", "")
        header = f"[{prefix}: {fname} | Type: {label}]"
        summary = (rd.get("ai_extracted_summary") or "").strip()
        if summary:
            text = summary[:MAX_SUMMARY]
            return f"{header}\n{text}"
        raw = rd.get("parsed_text", "")
        if raw:
            text = raw[:MAX_RAW_FALLBACK]
            if len(raw) > MAX_RAW_FALLBACK:
                text += "\n[... document truncated, AI extraction not yet run ...]"
            return f"{header}\n{text}"
        return None

    ai_docs = get_project_documents_for_ai(project_id)
    ref_parts = []
    for rd in ai_docs:
        if not rd.get("parsed_text") and not rd.get("ai_extracted_summary"):
            continue
        rd_type = rd.get("doc_type", "other")
        if rd_type == "pdd" and doc_type == "mr" and not pdd_text:
            pdd_text = rd.get("parsed_text", "")
            continue
        part = _doc_text(rd)
        if part:
            ref_parts.append(part)

    if project.get("parent_project_id"):
        parent_ai_docs = get_project_documents_for_ai(project["parent_project_id"])
        for rd in parent_ai_docs:
            part = _doc_text(rd, prefix="Parent Project Document")
            if part:
                ref_parts.append(part)

    combined = "\n\n---\n\n".join(ref_parts) if ref_parts else None
    if combined and len(combined) > MAX_TOTAL_REF:
        combined = combined[:MAX_TOTAL_REF] + "\n[... remaining documents truncated ...]"
    reference_text = combined
    return pdd_text, reference_text


@router.post("/{project_id}/write")
def write_section(project_id: int, data: WriteSectionRequest, doc_type: str = "pdd"):
    from carbongpt.repository.store import (
        get_user_project, save_write_session
    )
    from carbongpt.core.ai_writer import generate_section_draft

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    intake = project.get("project_intake") or {}
    if not intake.get("_project_brief"):
        try:
            from carbongpt.core.ai_writer import generate_project_brief as _gen_brief
            from carbongpt.repository.store import update_user_project as _upd_proj
            brief_info = {
                "name": project["name"],
                "standard": project.get("standard", ""),
                "methodology": project.get("methodology"),
                "country": project.get("country"),
                "description": project.get("description"),
                "project_intake": intake,
            }
            brief_chunks = None
            try:
                from carbongpt.core.project_doc_index import search_project_chunks as _sp, get_project_chunk_count as _gc
                if _gc(project_id) > 0:
                    brief_chunks = _sp(project_id, f"{project.get('name', '')} {project.get('methodology', '')} overview")
            except Exception:
                pass
            brief = _gen_brief(brief_info, brief_chunks)
            if brief:
                intake["_project_brief"] = brief
                _upd_proj(project_id, project_intake=intake)
                logger.info("Auto-generated project brief for project %s", project_id)
        except Exception as e:
            logger.warning("Auto project brief generation failed: %s", e)

    project_info = {
        "id": project["id"],
        "name": project["name"],
        "methodology": project.get("methodology"),
        "country": project.get("country"),
        "description": project.get("description"),
        "project_intake": intake,
        "project_settings": project.get("project_settings") or {},
        "selected_scenario_id": project.get("selected_scenario_id"),
    }

    pdd_text, reference_text = _gather_ai_context(project_id, project, doc_type)

    project_doc_context = None
    try:
        from carbongpt.core.project_doc_index import (
            search_project_chunks, format_project_context_for_prompt,
            build_section_search_query, get_project_chunk_count,
        )
        if get_project_chunk_count(project_id) > 0:
            from carbongpt.core.ai_writer import get_sections_for_doc_type as _get_secs
            all_secs = _get_secs(project["standard"], doc_type) or []
            must_include_items = []
            for s in all_secs:
                if s["id"] == data.section_id:
                    must_include_items = s.get("must_include", [])
                    break
            search_query = build_section_search_query(
                section_title=next((s["title"] for s in all_secs if s["id"] == data.section_id), data.section_id),
                must_include_items=must_include_items,
                project_info=project_info,
            )
            chunks = search_project_chunks(project_id, search_query)
            if chunks:
                project_doc_context = format_project_context_for_prompt(chunks)
                reference_text = None
    except Exception as e:
        logger.warning("Project doc RAG failed, falling back to concatenation: %s", e)

    try:
        generated = generate_section_draft(
            standard=project["standard"],
            project_doc_type=doc_type,
            section_id=data.section_id,
            project_info=project_info,
            existing_pdd_text=pdd_text,
            reference_docs_text=reference_text or project_doc_context,
            user_instructions=data.user_instructions,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error("AI writing failed: %s", e)
        raise HTTPException(status_code=500, detail=f"AI writing failed: {e}")

    from carbongpt.core.ai_writer import get_sections_for_doc_type
    sections = get_sections_for_doc_type(project["standard"], doc_type) or []
    section_title = ""
    for s in sections:
        if s["id"] == data.section_id:
            section_title = s["title"]
            break

    session_id = save_write_session(
        project_id=project_id,
        doc_type=doc_type,
        section_id=data.section_id,
        section_title=section_title,
        generated_text=generated,
    )

    validation = None
    try:
        from carbongpt.core.ai_writer import validate_section_output, get_sections_for_doc_type as _get_secs2
        all_secs2 = _get_secs2(project["standard"], doc_type) or []
        guide_section = None
        for s in all_secs2:
            if s["id"] == data.section_id:
                guide_section = s
                break
        if guide_section:
            validation = validate_section_output(generated, guide_section)
    except Exception as e:
        logger.warning("Output validation failed: %s", e)

    return {
        "session_id": session_id,
        "section_id": data.section_id,
        "section_title": section_title,
        "generated_text": generated,
        "validation": validation,
    }


@router.post("/{project_id}/write-all")
def write_all_sections(project_id: int, data: WriteAllRequest, doc_type: str = "pdd"):
    from carbongpt.repository.store import (
        get_user_project, save_write_session
    )
    from carbongpt.core.ai_writer import generate_full_document

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    project_info = {
        "id": project["id"],
        "name": project["name"],
        "methodology": project.get("methodology"),
        "country": project.get("country"),
        "description": project.get("description"),
        "project_intake": project.get("project_intake") or {},
        "project_settings": project.get("project_settings") or {},
        "selected_scenario_id": project.get("selected_scenario_id"),
    }

    pdd_text, reference_text = _gather_ai_context(project_id, project, doc_type)

    try:
        results = generate_full_document(
            standard=project["standard"],
            project_doc_type=doc_type,
            project_info=project_info,
            existing_pdd_text=pdd_text,
            reference_docs_text=reference_text,
            user_instructions=data.user_instructions,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error("AI full document generation failed: %s", e)
        raise HTTPException(status_code=500, detail=f"AI full document generation failed: {e}")

    for r in results:
        if r["status"] == "success" and r["generated_text"]:
            save_write_session(
                project_id=project_id,
                doc_type=doc_type,
                section_id=r["section_id"],
                section_title=r["section_title"],
                generated_text=r["generated_text"],
            )

    return {
        "sections": results,
        "total": len(results),
        "success_count": sum(1 for r in results if r["status"] == "success"),
    }


@router.patch("/{project_id}/section-text")
def update_section_text(project_id: int, data: UpdateSectionTextRequest):
    from carbongpt.repository.store import get_user_project, save_write_session
    from carbongpt.core.ai_writer import get_sections_for_doc_type

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    sections = get_sections_for_doc_type(project["standard"], data.doc_type) or []
    section_title = ""
    for s in sections:
        if s["id"] == data.section_id:
            section_title = s["title"]
            break

    session_id = save_write_session(
        project_id=project_id,
        doc_type=data.doc_type,
        section_id=data.section_id,
        section_title=section_title,
        generated_text=data.text,
        user_text=data.text,
    )

    return {"session_id": session_id, "section_id": data.section_id, "saved": True}


@router.post("/{project_id}/generate-brief")
def generate_brief_endpoint(project_id: int):
    from carbongpt.repository.store import get_user_project, update_user_project
    from carbongpt.core.ai_writer import generate_project_brief

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    project_info = {
        "name": project["name"],
        "standard": project.get("standard", ""),
        "methodology": project.get("methodology"),
        "country": project.get("country"),
        "description": project.get("description"),
        "project_intake": project.get("project_intake") or {},
    }

    doc_chunks = None
    try:
        from carbongpt.core.project_doc_index import search_project_chunks, get_project_chunk_count
        if get_project_chunk_count(project_id) > 0:
            doc_chunks = search_project_chunks(
                project_id,
                f"{project.get('name', '')} {project.get('methodology', '')} project overview technology",
            )
    except Exception:
        pass

    try:
        brief = generate_project_brief(project_info, doc_chunks)
    except Exception as e:
        logger.error("Project brief generation failed: %s", e)
        raise HTTPException(status_code=500, detail=f"Brief generation failed: {e}")

    intake = project.get("project_intake") or {}
    intake["_project_brief"] = brief
    update_user_project(project_id, project_intake=intake)

    return {"brief": brief}


@router.post("/{project_id}/validate-consistency")
def validate_consistency_endpoint(project_id: int, doc_type: str = "pdd"):
    from carbongpt.repository.store import get_user_project, get_write_sessions
    from carbongpt.core.ai_writer import validate_document_consistency

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    sessions = get_write_sessions(project_id, doc_type)
    if not sessions:
        return {"error": "No drafted sections found. Write sections first."}

    sections_dict = {}
    for s in sessions:
        sections_dict[s["section_id"]] = {
            "section_title": s.get("section_title", s["section_id"]),
            "generated_text": s.get("user_text") or s.get("generated_text", ""),
        }

    if len(sections_dict) < 2:
        return {"error": "Need at least 2 sections for consistency check."}

    project_info = {
        "name": project["name"],
        "methodology": project.get("methodology"),
        "country": project.get("country"),
    }

    try:
        result = validate_document_consistency(sections_dict, project_info)
    except Exception as e:
        logger.error("Consistency validation failed: %s", e)
        raise HTTPException(status_code=500, detail=f"Consistency validation failed: {e}")
    return result


@router.post("/{project_id}/explain")
def explain_section_endpoint(project_id: int, data: ExplainSectionRequest, doc_type: str = "pdd"):
    from carbongpt.repository.store import get_user_project
    from carbongpt.core.ai_writer import explain_section

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    try:
        explanation = explain_section(
            standard=project["standard"],
            project_doc_type=doc_type,
            section_id=data.section_id,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error("Explain section failed: %s", e)
        raise HTTPException(status_code=500, detail=f"Explain failed: {e}")

    return {"section_id": data.section_id, "explanation": explanation}


@router.post("/{project_id}/review/{doc_id}")
def review_document(project_id: int, doc_id: int):
    from carbongpt.repository.store import (
        get_user_project, get_project_document,
        get_project_documents_by_type, update_project_document
    )
    from carbongpt.core.ai_writer import review_with_context

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    doc = get_project_document(doc_id)
    if not doc or doc["project_id"] != project_id:
        raise HTTPException(status_code=404, detail="Document not found.")

    if not doc.get("parsed_text"):
        raise HTTPException(status_code=400, detail="Document has not been parsed. Please upload a DOCX or PDF file.")

    project_info = {
        "name": project["name"],
        "methodology": project.get("methodology"),
        "country": project.get("country"),
        "description": project.get("description"),
        "project_intake": project.get("project_intake") or {},
        "project_settings": project.get("project_settings") or {},
    }

    pdd_text, reference_text = _gather_ai_context(project_id, project, doc["doc_type"])

    try:
        review_result = review_with_context(
            standard=project["standard"],
            project_doc_type=doc["doc_type"],
            document_text=doc["parsed_text"],
            project_info=project_info,
            pdd_text=pdd_text,
            reference_texts=reference_text,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error("Review failed: %s", e)
        raise HTTPException(status_code=500, detail=f"Review failed: {e}")

    from psycopg2.extras import Json
    update_project_document(
        doc_id,
        review_result=Json(review_result),
        status="reviewed",
    )

    return review_result


@router.post("/{project_id}/review-draft")
def review_draft(project_id: int, doc_type: str = "pdd"):
    from carbongpt.repository.store import get_user_project, get_write_sessions
    from carbongpt.core.ai_writer import review_with_context

    valid_doc_types = {"pdd", "mr", "poa_dd", "vpa_dd", "valver"}
    if doc_type not in valid_doc_types:
        raise HTTPException(status_code=400, detail=f"Invalid doc_type: {doc_type}. Must be one of: {', '.join(sorted(valid_doc_types))}")

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    sessions = get_write_sessions(project_id, doc_type=doc_type)
    if not sessions:
        raise HTTPException(status_code=400, detail="No drafted sections found for this document type.")

    draft_parts = []
    for sess in sessions:
        text = sess.get("user_text") or sess.get("generated_text") or ""
        if text.strip():
            section_id = sess.get("section_id", "")
            section_title = sess.get("section_title", "")
            draft_parts.append(f"## {section_id} {section_title}\n\n{text}")

    if not draft_parts:
        raise HTTPException(status_code=400, detail="All drafted sections are empty.")

    document_text = "\n\n---\n\n".join(draft_parts)

    project_info = {
        "name": project["name"],
        "methodology": project.get("methodology"),
        "country": project.get("country"),
        "description": project.get("description"),
        "project_intake": project.get("project_intake") or {},
        "project_settings": project.get("project_settings") or {},
    }

    pdd_text, reference_text = _gather_ai_context(project_id, project, doc_type)

    try:
        review_result = review_with_context(
            standard=project["standard"],
            project_doc_type=doc_type,
            document_text=document_text,
            project_info=project_info,
            pdd_text=pdd_text,
            reference_texts=reference_text,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error("Draft review failed: %s", e)
        raise HTTPException(status_code=500, detail=f"Draft review failed: {e}")

    return review_result


@router.post("/{project_id}/respond-to-finding")
def respond_to_finding(project_id: int, data: dict):
    from carbongpt.repository.store import get_user_project, get_write_sessions
    from carbongpt.core.ai_writer import _call_openai

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    finding_text = data.get("finding_text", "")
    finding_type = data.get("finding_type", "CL")
    pdd_section = data.get("pdd_section", "")
    doc_type = data.get("doc_type", "pdd")

    if not finding_text:
        raise HTTPException(status_code=400, detail="finding_text is required")

    project_info = {
        "name": project["name"],
        "methodology": project.get("methodology"),
        "country": project.get("country"),
        "description": project.get("description"),
        "project_intake": project.get("project_intake") or {},
        "project_settings": project.get("project_settings") or {},
    }

    pdd_text, reference_text = _gather_ai_context(project_id, project, doc_type)

    relevant_section_text = ""
    if pdd_section:
        sessions = get_write_sessions(project_id, doc_type=doc_type)
        for sess in sessions:
            if pdd_section.lower() in (sess.get("section_id", "") or "").lower():
                relevant_section_text = sess.get("user_text") or sess.get("generated_text") or ""
                break

    from carbongpt.core.ai_writer import STANDARD_LABELS, DOC_TYPE_LABELS, get_guide_doc_type
    std_label = STANDARD_LABELS.get(project["standard"], project["standard"])

    system_prompt = (
        f"You are an expert carbon project developer responding to a {finding_type} "
        f"(finding) raised by a VVB or standard body ({std_label}) during validation/verification.\n\n"
        "RULES:\n"
        "- Draft a professional, detailed response that directly addresses the finding.\n"
        "- Reference specific project data, methodology clauses, and evidence.\n"
        "- If the PDD or document needs to be updated, clearly state what changes are needed.\n"
        "- Follow the standard response format: acknowledge the issue, explain the resolution, cite evidence.\n"
        "- Be thorough but concise. VVBs appreciate clear, well-structured responses.\n"
        "- Format your response as JSON with keys: 'response_text', 'pdd_updates_needed' (array of {section, change_description}), "
        "'evidence_to_provide' (array of strings), 'response_approach' (one of: pdd_update, clarification, evidence_provided, calculation_corrected, methodology_reference)"
    )

    user_prompt = f"## Finding to respond to:\n"
    user_prompt += f"**Type:** {finding_type}\n"
    user_prompt += f"**PDD Section:** {pdd_section or 'Not specified'}\n"
    user_prompt += f"**Finding:**\n{finding_text}\n\n"

    user_prompt += f"### Project Information:\n"
    user_prompt += f"- Project: {project_info.get('name', 'Unknown')}\n"
    user_prompt += f"- Standard: {std_label}\n"
    user_prompt += f"- Methodology: {project_info.get('methodology', 'Not specified')}\n"
    user_prompt += f"- Country: {project_info.get('country', 'Not specified')}\n\n"

    if relevant_section_text:
        user_prompt += (
            f"### Current content of section {pdd_section}:\n"
            f'"""\n{relevant_section_text[:4000]}\n"""\n\n'
        )

    if pdd_text:
        user_prompt += (
            "### PDD content for reference:\n"
            f'"""\n{pdd_text[:4000]}\n"""\n\n'
        )

    findings_context = ""
    try:
        methodology = project_info.get("methodology", "")
        if methodology:
            from carbongpt.repository.store import get_findings_by_methodology
            similar = get_findings_by_methodology(methodology, limit=20)
            relevant_similar = [
                f for f in similar
                if f.get("resolution") and (
                    (f.get("topic") or "").lower() in finding_text.lower()
                    or (f.get("pdd_section") or "").lower() == (pdd_section or "").lower()
                )
            ][:5]
            if relevant_similar:
                findings_context = "### How similar findings were resolved on other projects:\n"
                for sf in relevant_similar:
                    findings_context += f"- [{sf['finding_type']}] {sf.get('topic', '')}: {sf.get('description', '')[:200]}\n"
                    findings_context += f"  Resolution: {sf.get('resolution', '')[:200]}\n"
                    findings_context += f"  Approach: {sf.get('resolution_approach', '')}\n\n"
    except Exception as e:
        logger.warning("Failed to get similar findings: %s", e)

    if findings_context:
        user_prompt += findings_context + "\n"

    user_prompt += "Draft a professional response to this finding."

    try:
        import json
        from carbongpt.core.openai_client import call_openai
        content = call_openai(system_prompt, user_prompt, max_tokens=3000, temperature=0.3)
        result = json.loads(content)
        return result
    except Exception as e:
        logger.error("Finding response generation failed: %s", e)
        raise HTTPException(status_code=500, detail=f"AI response generation failed: {e}")


@router.post("/{project_id}/parse-findings-document")
def parse_findings_document(
    project_id: int,
    file: UploadFile = File(...),
):
    from carbongpt.repository.store import get_user_project

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ("pdf", "docx"):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")

    MAX_FILE_SIZE = 20 * 1024 * 1024
    import tempfile
    content = file.file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)} MB.")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}")
    tmp.write(content)
    tmp.close()

    parsed_text = ""
    try:
        if ext == "pdf":
            import pdfplumber
            text_parts = []
            with pdfplumber.open(tmp.name) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
            parsed_text = "\n".join(text_parts)
        elif ext == "docx":
            from carbongpt.tools.parse_docx import parse_docx
            result = parse_docx(tmp.name)
            sections = result.get("sections", {})
            parsed_text = "\n\n".join(
                f"### {k}\n{v}" if k != "__PREAMBLE__" else v
                for k, v in sections.items() if v
            )
    except Exception as e:
        logger.error("Failed to parse findings document: %s", e)
        raise HTTPException(status_code=500, detail=f"Failed to parse document: {e}")
    finally:
        os.unlink(tmp.name)

    if not parsed_text or len(parsed_text.strip()) < 50:
        raise HTTPException(status_code=400, detail="Could not extract meaningful text from the document.")

    max_chunk_chars = 12000
    chunks = []
    words = parsed_text.split()
    current_chunk = []
    current_len = 0
    for word in words:
        current_chunk.append(word)
        current_len += len(word) + 1
        if current_len >= max_chunk_chars:
            chunks.append(" ".join(current_chunk))
            current_chunk = []
            current_len = 0
    if current_chunk:
        chunks.append(" ".join(current_chunk))

    extraction_prompt = (
        "You are an expert carbon credit auditor. Extract ALL individual findings "
        "(CARs, CLs, FARs, observations, PRR comments) from the document text below.\n\n"
        "For each finding, provide:\n"
        "- finding_type: 'CAR' | 'CL' | 'FAR' | 'observation' | 'prr_comment'\n"
        "- finding_id: the original ID from the document (e.g., 'CAR #1', 'CL 05')\n"
        "- pdd_section: which PDD/MR section this relates to\n"
        "- description: the VVB's or reviewer's question/concern (the actual finding text)\n"
        "- topic: short categorization (e.g., 'baseline scenario', 'monitoring parameters')\n\n"
        "Return JSON: {\"findings\": [...]}. If no findings found, return {\"findings\": []}.\n"
        "Be thorough - extract EVERY finding including minor clarifications.\n\n"
        "Document text:\n"
    )

    all_findings = []
    failed_chunks = []
    import json
    from carbongpt.core.openai_client import call_openai

    for i, chunk in enumerate(chunks):
        try:
            content = call_openai(extraction_prompt, chunk, max_tokens=4000, temperature=0.1)
            result = json.loads(content)
            chunk_findings = result.get("findings", [])
            all_findings.extend(chunk_findings)
            logger.info("Chunk %d/%d: extracted %d findings", i + 1, len(chunks), len(chunk_findings))
        except Exception as e:
            logger.warning("Findings extraction failed for chunk %d: %s", i + 1, e)
            failed_chunks.append({"chunk": i + 1, "error": str(e)})

    if failed_chunks and not all_findings:
        raise HTTPException(
            status_code=500,
            detail=f"AI extraction failed for all {len(chunks)} document sections. First error: {failed_chunks[0]['error']}"
        )

    seen = set()
    unique_findings = []
    for f in all_findings:
        key = (f.get("finding_id", ""), f.get("description", "")[:80])
        if key not in seen:
            seen.add(key)
            unique_findings.append(f)

    return {
        "findings": unique_findings,
        "total": len(unique_findings),
        "document_name": file.filename,
        "text_length": len(parsed_text),
        "chunks_processed": len(chunks),
        "chunks_failed": len(failed_chunks),
        "warnings": [f"Chunk {fc['chunk']} failed: {fc['error'][:100]}" for fc in failed_chunks] if failed_chunks else [],
    }


@router.post("/{project_id}/batch-respond-to-findings")
def batch_respond_to_findings(project_id: int, data: dict):
    from carbongpt.repository.store import get_user_project, get_write_sessions

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    findings = data.get("findings", [])
    doc_type = data.get("doc_type", "pdd")

    if not findings:
        raise HTTPException(status_code=400, detail="No findings provided.")

    project_info = {
        "name": project["name"],
        "methodology": project.get("methodology"),
        "country": project.get("country"),
        "description": project.get("description"),
        "project_intake": project.get("project_intake") or {},
    }

    pdd_text, reference_text = _gather_ai_context(project_id, project, doc_type)

    sessions = get_write_sessions(project_id, doc_type=doc_type)
    session_map = {}
    for sess in sessions:
        sid = (sess.get("section_id") or "").lower()
        session_map[sid] = sess.get("user_text") or sess.get("generated_text") or ""

    from carbongpt.core.ai_writer import STANDARD_LABELS
    std_label = STANDARD_LABELS.get(project["standard"], project["standard"])

    findings_context = ""
    try:
        methodology = project_info.get("methodology", "")
        if methodology:
            from carbongpt.repository.store import get_findings_by_methodology
            similar = get_findings_by_methodology(methodology, limit=30)
            if similar:
                findings_context = "### How similar findings were resolved on other projects:\n"
                for sf in similar[:10]:
                    if sf.get("resolution"):
                        findings_context += f"- [{sf['finding_type']}] {sf.get('topic', '')}: {sf.get('description', '')[:150]}\n"
                        findings_context += f"  Resolution: {sf.get('resolution', '')[:150]}\n\n"
    except Exception:
        pass

    responses = []
    import json
    from carbongpt.core.openai_client import call_openai

    for idx, finding in enumerate(findings):
        finding_text = finding.get("description", "")
        finding_type = finding.get("finding_type", "CL")
        finding_id = finding.get("finding_id", f"Finding {idx + 1}")
        pdd_section = finding.get("pdd_section", "")

        relevant_section_text = ""
        if pdd_section:
            for sid, text in session_map.items():
                if pdd_section.lower() in sid:
                    relevant_section_text = text
                    break

        system_prompt = (
            f"You are an expert carbon project developer responding to a {finding_type} "
            f"raised by a VVB or standard body ({std_label}).\n\n"
            "RULES:\n"
            "- Draft a professional, detailed response that directly addresses the finding.\n"
            "- Reference specific project data, methodology clauses, and evidence.\n"
            "- If the PDD/document needs updating, clearly state what changes are needed.\n"
            "- Be thorough but concise.\n"
            "- Format response as JSON with keys: 'response_text', 'pdd_updates_needed' (array of {section, change_description}), "
            "'evidence_to_provide' (array of strings), 'response_approach' (one of: pdd_update, clarification, evidence_provided, calculation_corrected, methodology_reference)"
        )

        user_prompt = f"## Finding: {finding_id}\n"
        user_prompt += f"**Type:** {finding_type}\n"
        user_prompt += f"**PDD Section:** {pdd_section or 'Not specified'}\n"
        user_prompt += f"**Finding:**\n{finding_text}\n\n"
        user_prompt += f"### Project Information:\n"
        user_prompt += f"- Project: {project_info.get('name', 'Unknown')}\n"
        user_prompt += f"- Standard: {std_label}\n"
        user_prompt += f"- Methodology: {project_info.get('methodology', 'Not specified')}\n"
        user_prompt += f"- Country: {project_info.get('country', 'Not specified')}\n\n"

        if relevant_section_text:
            user_prompt += f"### Current content of section {pdd_section}:\n\"\"\"\n{relevant_section_text[:3000]}\n\"\"\"\n\n"
        if pdd_text:
            user_prompt += f"### PDD content:\n\"\"\"\n{pdd_text[:3000]}\n\"\"\"\n\n"
        if findings_context:
            user_prompt += findings_context + "\n"

        user_prompt += "Draft a professional response."

        try:
            content = call_openai(system_prompt, user_prompt, max_tokens=2500, temperature=0.3)
            ai_result = json.loads(content)
            responses.append({
                "finding_id": finding_id,
                "finding_type": finding_type,
                "finding_text": finding_text,
                "pdd_section": pdd_section,
                "topic": finding.get("topic", ""),
                "status": "success",
                **ai_result,
            })
            logger.info("Generated response for %s (%d/%d)", finding_id, idx + 1, len(findings))
        except Exception as e:
            logger.warning("Failed to generate response for %s: %s", finding_id, e)
            responses.append({
                "finding_id": finding_id,
                "finding_type": finding_type,
                "finding_text": finding_text,
                "pdd_section": pdd_section,
                "topic": finding.get("topic", ""),
                "status": "error",
                "error": str(e),
                "response_text": "",
                "pdd_updates_needed": [],
                "evidence_to_provide": [],
                "response_approach": "",
            })

    return {
        "responses": responses,
        "total": len(findings),
        "successful": sum(1 for r in responses if r["status"] == "success"),
        "failed": sum(1 for r in responses if r["status"] == "error"),
    }


@router.get("/{project_id}/write-sessions")
def get_write_sessions_endpoint(project_id: int, doc_type: str = None):
    from carbongpt.repository.store import get_write_sessions
    return get_write_sessions(project_id, doc_type=doc_type)


@router.get("/{project_id}/methodology-data")
def get_methodology_data_endpoint(project_id: int):
    from carbongpt.repository.store import get_user_project, get_parsed_methodology

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    methodology = project.get("methodology")
    if not methodology:
        return {"status": "no_methodology", "parsed": None}

    cached = get_parsed_methodology(methodology)
    if cached and cached.get("parsed_data"):
        data = cached["parsed_data"]
        if isinstance(data, str):
            import json as _json
            data = _json.loads(data)
        try:
            from carbongpt.core.tool_defaults import enrich_methodology_parameters
            country = project.get("country", "")
            data = enrich_methodology_parameters(data, methodology, country=country)
        except Exception as e:
            logger.warning("TOOL33 parameter enrichment failed: %s", e)
        return {"status": "ready", "parsed": data, "parsed_at": str(cached.get("parsed_at", ""))}

    return {"status": "not_parsed", "parsed": None}


@router.post("/{project_id}/parse-methodology")
def parse_methodology_endpoint(project_id: int, data: ParseMethodologyRequest):
    from carbongpt.repository.store import get_user_project
    from carbongpt.core.methodology_parser import get_methodology_sections

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    force = data.force if hasattr(data, 'force') else False

    try:
        from carbongpt.core.methodology_kb import build_methodology_knowledge
        sections_data = get_methodology_sections(data.methodology_code)
        if not sections_data:
            raise ValueError(f"No methodology document found for '{data.methodology_code}'")
        result = build_methodology_knowledge(sections_data["doc_id"], data.methodology_code, force=force)
        from carbongpt.core.methodology_kb import get_knowledge_as_parsed_format
        parsed = get_knowledge_as_parsed_format(data.methodology_code)
        if parsed:
            return parsed
        return result
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error("Methodology KB build failed, falling back to legacy parser: %s", e)
        from carbongpt.core.methodology_parser import parse_methodology_and_save
        try:
            parsed = parse_methodology_and_save(data.methodology_code, force=force)
            return parsed
        except Exception as e2:
            logger.error("Legacy parse also failed: %s", e2)
            raise HTTPException(status_code=500, detail=f"Failed to parse methodology: {e2}")


@router.post("/{project_id}/calculate")
def run_calculation_endpoint(project_id: int, data: RunCalculationRequest):
    from carbongpt.repository.store import get_user_project
    from carbongpt.core.methodology_parser import get_or_parse_methodology
    from carbongpt.core.calculation_engine import run_calculation

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    methodology = project.get("methodology")
    if not methodology:
        raise HTTPException(status_code=400, detail="Project has no methodology assigned.")

    try:
        parsed = get_or_parse_methodology(methodology)
    except Exception as e:
        logger.error("Methodology parse failed for calc: %s", e)
        raise HTTPException(status_code=500, detail=f"Failed to parse methodology: {e}")

    project_info = {
        "name": project["name"],
        "standard": project.get("standard"),
        "country": project.get("country"),
        "description": project.get("description"),
    }

    try:
        calc_result = run_calculation(
            parsed_methodology=parsed,
            user_inputs=data.user_inputs,
            method_id=data.method_id,
            crediting_years=data.crediting_years,
            project_info=project_info,
        )
    except Exception as e:
        logger.error("Calculation failed: %s", e)
        raise HTTPException(status_code=500, detail=f"Calculation failed: {e}")

    return calc_result


@router.post("/{project_id}/export-calculation")
def export_calculation_endpoint(project_id: int, data: ExportCalculationRequest):
    from carbongpt.repository.store import get_user_project
    from carbongpt.core.er_excel import generate_er_workbook

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    try:
        # Prefer the selected/saved scenario; fall back to session payload
        calc_result = data.calculation_result or {}
        try:
            from carbongpt.core.er_simulator import get_selected_scenario
            selected = get_selected_scenario(project_id)
            if selected and isinstance(selected, dict) and "annual_results" in selected:
                calc_result = selected
        except Exception as _se:
            logger.debug("Could not load selected scenario, using session payload: %s", _se)

        buf = generate_er_workbook(
            project=project,
            doc_type=data.doc_type or "pdd",
            calc_result=calc_result,
            project_id=project_id,
        )
    except Exception as e:
        logger.error("ER workbook export failed: %s", e)
        raise HTTPException(status_code=500, detail=f"Export failed: {e}")

    doc_suffix = "ExPost" if data.doc_type in ("mr",) else "ExAnte"
    safe_name = project["name"].replace(" ", "_")[:30]
    filename = f"{safe_name}_ER_{doc_suffix}.xlsx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/{project_id}/generate-template")
def generate_template_endpoint(project_id: int, data: GenerateTemplateRequest):
    from carbongpt.repository.store import get_user_project, get_write_sessions
    from carbongpt.core.ai_writer import get_sections_for_doc_type
    from carbongpt.core.doc_exporter import generate_filled_template

    project = get_user_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found.")

    intake = {}
    if project.get("project_intake"):
        try:
            import json
            raw = project["project_intake"]
            intake = json.loads(raw) if isinstance(raw, str) else (raw or {})
        except Exception:
            intake = {}

    project_info = {
        "name": project["name"],
        "standard": project.get("standard"),
        "methodology": project.get("methodology"),
        "country": project.get("country"),
        "description": project.get("description"),
        "doc_type": data.doc_type,
        "intake": intake,
    }

    sections = get_sections_for_doc_type(project["standard"], data.doc_type)
    if not sections:
        raise HTTPException(status_code=400,
                            detail=f"No template guide found for {project['standard']}/{data.doc_type}")

    write_sessions = get_write_sessions(project_id, doc_type=data.doc_type)
    session_map = {ws["section_id"]: ws for ws in write_sessions}

    generated_sections = []
    for sec in sections:
        sid = sec["id"]
        content = ""
        ws = session_map.get(sid)
        if ws:
            content = ws.get("user_text") or ws.get("generated_text") or ""
        generated_sections.append({
            "section_id": sid,
            "title": sec.get("title", ""),
            "content": content,
        })

    calc_result = None
    if data.include_calculations and data.calculation_result:
        calc_result = data.calculation_result

    try:
        buf = generate_filled_template(project_info, generated_sections, calc_result=calc_result)
    except Exception as e:
        logger.error("Template generation failed: %s", e)
        raise HTTPException(status_code=500, detail=f"Template generation failed: {e}")

    safe_name = project["name"].replace(" ", "_")[:30]
    doc_type_label = data.doc_type.upper().replace("_", "-")
    filename = f"{safe_name}_{doc_type_label}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


class ChatMessage(BaseModel):
    message: str
    project_id: int | None = None
    history: list[dict] = []


@router.post("/chat")
def chat_with_ai(body: ChatMessage):
    from carbongpt.core.copilot import process_copilot_message

    project_context = ""
    if body.project_id:
        from carbongpt.repository.store import get_user_project, get_project_documents_for_ai
        project = get_user_project(body.project_id)
        if project:
            intake = project.get("project_intake") or {}
            context_parts = [
                f"Project: {project['name']}",
                f"Standard: {project.get('standard', '')}",
                f"Type: {project.get('project_type', '')}",
                f"Methodology: {project.get('methodology', '')}",
                f"Country: {project.get('country', '')}",
                f"Status: {project.get('status', '')}",
            ]

            if intake:
                filtered = {k: v for k, v in intake.items() if not k.startswith("_") and v}
                if filtered:
                    import json
                    context_parts.append(f"Project Setup Data: {json.dumps(filtered, default=str)[:3000]}")

            docs = get_project_documents_for_ai(body.project_id)
            if docs:
                doc_summaries = []
                for doc in docs[:5]:
                    summary = doc.get("ai_extracted_summary", "")
                    if summary:
                        doc_summaries.append(f"[{doc['file_name']}]: {summary[:800]}")
                if doc_summaries:
                    context_parts.append("Document Intelligence:\n" + "\n".join(doc_summaries))

            from carbongpt.repository.store import get_write_sessions
            sessions = get_write_sessions(body.project_id)
            if sessions:
                drafted = [s for s in sessions if s.get("generated_text") or s.get("user_text")]
                if drafted:
                    draft_info = [f"Section {s['section_id']} ({s.get('section_title', '')}): {(s.get('user_text') or s.get('generated_text', ''))[:200]}..." for s in drafted[:10]]
                    context_parts.append("Drafted Sections:\n" + "\n".join(draft_info))

            # ── Step 7: Carbon Intelligence market context injection ──────────
            # Fetch registry-level market data for the project's country and
            # methodology family.  This gives the AI factual grounding on
            # comparable projects without querying the DB directly.
            try:
                from carbongpt.repository.country_normalizer import resolve_country_iso
                from carbongpt.repository.store import get_market_intelligence_context

                ci_parts: list[str] = []
                country_iso = None
                raw_country = project.get("country") or ""
                if raw_country:
                    country_iso = resolve_country_iso(raw_country)
                methodology = project.get("methodology") or ""

                if country_iso or methodology:
                    ci_ctx = get_market_intelligence_context(
                        country_iso=country_iso or None,
                        methodology_family=methodology[:50] if methodology else None,
                    )
                    raw_prompt = ci_ctx.get("raw_for_prompt", "")
                    if raw_prompt:
                        ci_parts.append(raw_prompt)

                if ci_parts:
                    context_parts.append(
                        "\n--- Carbon Intelligence Market Context ---\n"
                        "(Registry-declared estimates, not verified issued credits)\n"
                        + "\n".join(ci_parts)
                    )
            except Exception as _ci_err:
                logger.debug("CI context injection skipped: %s", _ci_err)

            project_context = "\n".join(context_parts)

    try:
        result = process_copilot_message(
            message=body.message,
            project_id=body.project_id,
            history=body.history,
            project_context=project_context,
        )
        return result
    except Exception as e:
        logger.error("Chat error: %s", e)
        err_str = str(e)
        if "429" in err_str or "rate" in err_str.lower():
            raise HTTPException(status_code=429, detail="Rate limit reached. Please wait and try again.")
        raise HTTPException(status_code=500, detail="Chat failed. Please try again.")


class ResearchRunRequest(BaseModel):
    doc_type: str = "pdd"
    max_gaps: int = Field(default=20, le=50)


class ResearchConfirmRequest(BaseModel):
    result_id: int


@router.post("/{project_id}/research/analyze-gaps")
def analyze_gaps_endpoint(project_id: int, doc_type: str = "pdd"):
    from carbongpt.core.research_orchestrator import analyze_gaps
    try:
        result = analyze_gaps(project_id, doc_type)
    except Exception as e:
        logger.error("Gap analysis failed: %s", e)
        raise HTTPException(status_code=500, detail=f"Gap analysis failed: {e}")
    return result


@router.post("/{project_id}/research/run")
def run_research_endpoint(project_id: int, data: ResearchRunRequest):
    from carbongpt.core.research_orchestrator import run_research_session
    try:
        result = run_research_session(project_id, data.doc_type, data.max_gaps)
    except Exception as e:
        logger.error("Research session failed: %s", e)
        raise HTTPException(status_code=500, detail=f"Research session failed: {e}")
    return result


@router.post("/{project_id}/research/confirm")
def confirm_research_endpoint(project_id: int, data: ResearchConfirmRequest):
    from carbongpt.core.research_orchestrator import confirm_research_result
    try:
        result = confirm_research_result(data.result_id, project_id)
    except Exception as e:
        logger.error("Research confirm failed: %s", e)
        raise HTTPException(status_code=500, detail=f"Confirm failed: {e}")
    return result


@router.post("/{project_id}/research/reject")
def reject_research_endpoint(project_id: int, data: ResearchConfirmRequest):
    from carbongpt.core.research_orchestrator import reject_research_result
    try:
        result = reject_research_result(data.result_id, project_id)
    except Exception as e:
        logger.error("Research reject failed: %s", e)
        raise HTTPException(status_code=500, detail=f"Reject failed: {e}")
    return result


@router.get("/{project_id}/research/results")
def get_research_results_endpoint(project_id: int, status: str = None):
    from carbongpt.core.research_orchestrator import get_research_results
    try:
        results = get_research_results(project_id, status)
    except Exception as e:
        logger.error("Get research results failed: %s", e)
        raise HTTPException(status_code=500, detail=f"Failed to get results: {e}")
    return {"results": results}


@router.post("/{project_id}/parameters/initialize")
def initialize_parameters_endpoint(project_id: int):
    from carbongpt.core.parameter_engine import initialize_project_parameters
    try:
        result = initialize_project_parameters(project_id)
    except Exception as e:
        logger.error("Parameter init failed: %s", e)
        raise HTTPException(status_code=500, detail=str(e))
    return result


@router.get("/{project_id}/parameters")
def get_parameters_endpoint(project_id: int, category: str = None):
    from carbongpt.core.parameter_engine import get_project_parameters
    try:
        params = get_project_parameters(project_id, category)
    except Exception as e:
        logger.error("Get parameters failed: %s", e)
        raise HTTPException(status_code=500, detail=str(e))
    return {"parameters": params}


@router.get("/{project_id}/parameters/summary")
def get_parameters_summary_endpoint(project_id: int):
    from carbongpt.core.parameter_engine import get_parameter_summary
    try:
        summary = get_parameter_summary(project_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return summary


class ParameterUpdate(BaseModel):
    value: str | None = None
    source_type: str | None = None
    source_reference: str | None = None
    notes: str | None = None


@router.put("/{project_id}/parameters/{param_key}")
def update_parameter_endpoint(project_id: int, param_key: str, data: ParameterUpdate):
    from carbongpt.core.parameter_engine import update_parameter
    try:
        result = update_parameter(project_id, param_key, data.value, data.source_type, data.source_reference, data.notes)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not result:
        raise HTTPException(status_code=404, detail="Parameter not found")
    return result


@router.post("/{project_id}/parameters/validate")
def validate_parameters_endpoint(project_id: int):
    from carbongpt.core.parameter_engine import validate_all_parameters
    try:
        result = validate_all_parameters(project_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return result


@router.post("/{project_id}/er-scenarios/calculate")
def calculate_er_endpoint(project_id: int):
    from carbongpt.core.er_simulator import run_scenario
    try:
        result = run_scenario(project_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return result


class ScenarioSave(BaseModel):
    name: str
    description: str = ""
    parameter_overrides: dict = {}
    carbon_price: float | None = None
    price_escalation: float = 0
    developer_share: float = 100
    buffer_pool: float = 0
    admin_fee: float = 0
    is_baseline: bool = False
    scenario_purpose: str = "exploratory"


@router.post("/{project_id}/er-scenarios")
def save_scenario_endpoint(project_id: int, data: ScenarioSave):
    from carbongpt.core.er_simulator import save_scenario
    try:
        result = save_scenario(
            project_id, data.name, data.description, data.parameter_overrides,
            data.carbon_price, data.price_escalation, data.developer_share,
            data.buffer_pool, data.admin_fee, data.is_baseline,
            scenario_purpose=data.scenario_purpose,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return result


@router.get("/{project_id}/er-scenarios")
def get_scenarios_endpoint(project_id: int):
    from carbongpt.core.er_simulator import get_scenarios, migrate_baseline_to_selected
    try:
        migrate_baseline_to_selected(project_id)
        scenarios = get_scenarios(project_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"scenarios": scenarios}


@router.get("/{project_id}/er-scenarios/selected")
def get_selected_scenario_endpoint(project_id: int):
    from carbongpt.core.er_simulator import get_selected_scenario
    try:
        result = get_selected_scenario(project_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if not result:
        return {"selected": None}
    return {"selected": result}


@router.post("/{project_id}/er-scenarios/{scenario_id}/select")
def select_scenario_endpoint(project_id: int, scenario_id: int):
    from carbongpt.core.er_simulator import select_scenario_for_drafting
    try:
        result = select_scenario_for_drafting(project_id, scenario_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if "error" in result:
        raise HTTPException(status_code=404, detail=result["error"])
    return result


@router.post("/{project_id}/er-scenarios/deselect")
def deselect_scenario_endpoint(project_id: int):
    from carbongpt.core.er_simulator import deselect_scenario
    try:
        result = deselect_scenario(project_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return result


class ScenarioPurposeUpdate(BaseModel):
    purpose: str


@router.patch("/{project_id}/er-scenarios/{scenario_id}/purpose")
def update_scenario_purpose_endpoint(project_id: int, scenario_id: int, data: ScenarioPurposeUpdate):
    from carbongpt.core.er_simulator import update_scenario_purpose
    try:
        result = update_scenario_purpose(project_id, scenario_id, data.purpose)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])
    return result


@router.get("/{project_id}/er-scenarios/compare")
def compare_scenarios_endpoint(project_id: int):
    from carbongpt.core.er_simulator import compare_scenarios
    try:
        result = compare_scenarios(project_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return result


@router.post("/{project_id}/er-scenarios/{param_key}/sensitivity")
def run_sensitivity_endpoint(project_id: int, param_key: str, variation: int = 20):
    from carbongpt.core.er_simulator import run_sensitivity
    try:
        result = run_sensitivity(project_id, param_key, variation_pct=variation)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return result


@router.get("/{project_id}/state")
def get_project_state_endpoint(project_id: int):
    from carbongpt.core.project_state import evaluate_project_state
    try:
        result = evaluate_project_state(project_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    if "error" in result:
        raise HTTPException(status_code=404, detail=result["error"])
    return result


@router.get("/{project_id}/lifecycle")
def get_lifecycle_endpoint(project_id: int):
    from carbongpt.core.lifecycle_manager import get_lifecycle
    try:
        result = get_lifecycle(project_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return result


@router.post("/{project_id}/lifecycle/initialize")
def initialize_lifecycle_endpoint(project_id: int):
    from carbongpt.core.lifecycle_manager import initialize_lifecycle
    try:
        result = initialize_lifecycle(project_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return result


@router.post("/{project_id}/lifecycle/advance")
def advance_stage_endpoint(project_id: int, to_stage: str = None):
    from carbongpt.core.lifecycle_manager import advance_stage
    try:
        result = advance_stage(project_id, to_stage)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return result


class TaskCreate(BaseModel):
    title: str
    stage: str | None = None
    task_type: str = "general"
    priority: str = "medium"
    due_date: str | None = None
    description: str | None = None


@router.get("/{project_id}/tasks")
def get_tasks_endpoint(project_id: int, stage: str = None, status: str = None):
    from carbongpt.core.lifecycle_manager import get_tasks
    try:
        tasks = get_tasks(project_id, stage, status)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"tasks": tasks}


@router.post("/{project_id}/tasks")
def add_task_endpoint(project_id: int, data: TaskCreate):
    from carbongpt.core.lifecycle_manager import add_task
    try:
        result = add_task(project_id, data.title, data.stage, data.task_type, data.priority, data.due_date, data.description)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return result


class TaskUpdate(BaseModel):
    status: str | None = None
    title: str | None = None
    due_date: str | None = None
    priority: str | None = None


@router.put("/tasks/{task_id}")
def update_task_endpoint(task_id: int, data: TaskUpdate):
    from carbongpt.core.lifecycle_manager import update_task
    try:
        result = update_task(task_id, data.status, data.title, data.due_date, data.priority)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return result


@router.get("/{project_id}/evidence")
def get_evidence_endpoint(project_id: int, target_type: str = None, target_id: str = None):
    from carbongpt.core.evidence_engine import get_evidence_links
    try:
        links = get_evidence_links(project_id, target_type, target_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"evidence": links}


@router.get("/{project_id}/evidence/completeness")
def get_evidence_completeness_endpoint(project_id: int):
    from carbongpt.core.evidence_engine import get_evidence_completeness
    try:
        result = get_evidence_completeness(project_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return result


@router.get("/{project_id}/evidence/citations")
def get_citations_endpoint(project_id: int):
    from carbongpt.core.evidence_engine import generate_citation_list
    try:
        citations = generate_citation_list(project_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"citations": citations}


@router.post("/{project_id}/documents/{doc_id}/extract-evidence")
def extract_evidence_endpoint(project_id: int, doc_id: int):
    from carbongpt.core.evidence_engine import extract_parameter_evidence
    try:
        result = extract_parameter_evidence(project_id, doc_id)
    except Exception as e:
        logger.error("Evidence extraction failed: %s", e)
        raise HTTPException(status_code=500, detail=str(e))
    if result.get("error"):
        raise HTTPException(status_code=400, detail=result["error"])
    return result


@router.get("/{project_id}/evidence/pending")
def get_pending_evidence_endpoint(project_id: int, doc_id: int = None):
    from carbongpt.core.evidence_engine import get_pending_evidence
    try:
        items = get_pending_evidence(project_id, doc_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"pending": items, "count": len(items)}


@router.post("/{project_id}/evidence/{link_id}/decide")
def decide_evidence_endpoint(project_id: int, link_id: int, body: dict):
    from carbongpt.core.evidence_engine import decide_evidence, decide_evidence_force
    decision = body.get("decision")
    force = body.get("force", False)
    if not decision:
        raise HTTPException(status_code=400, detail="decision is required")
    try:
        if force:
            result = decide_evidence_force(project_id, link_id, decision)
        else:
            result = decide_evidence(project_id, link_id, decision)
    except Exception as e:
        logger.error("Evidence decision failed: %s", e)
        raise HTTPException(status_code=500, detail=str(e))
    if result.get("error"):
        raise HTTPException(status_code=400, detail=result["error"])
    return result


@router.post("/{project_id}/audit-simulation")
def run_audit_simulation_endpoint(project_id: int):
    from carbongpt.core.audit_simulator import run_audit_simulation
    try:
        result = run_audit_simulation(project_id)
    except Exception as e:
        logger.error("Audit simulation failed: %s", e)
        raise HTTPException(status_code=500, detail=str(e))
    return result


@router.get("/{project_id}/audit-simulation/history")
def get_audit_history_endpoint(project_id: int):
    from carbongpt.core.audit_simulator import get_simulation_history
    try:
        history = get_simulation_history(project_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"history": history}


@router.get("/{project_id}/monitoring-tasks")
def get_monitoring_tasks_endpoint(project_id: int):
    from carbongpt.core.lifecycle_manager import get_monitoring_tasks
    try:
        tasks = get_monitoring_tasks(project_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"tasks": tasks}


@router.post("/{project_id}/monitoring-tasks/initialize")
def initialize_monitoring_endpoint(project_id: int):
    from carbongpt.core.lifecycle_manager import initialize_monitoring_tasks
    try:
        tasks = initialize_monitoring_tasks(project_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"tasks": tasks}


@router.get("/{project_id}/issuances")
def get_issuances_endpoint(project_id: int):
    from carbongpt.core.lifecycle_manager import get_issuances
    try:
        issuances = get_issuances(project_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return {"issuances": issuances}


class IssuanceCreate(BaseModel):
    vintage_year: int
    credits_requested: float | None = None
    credits_issued: float | None = None
    monitoring_period_start: str | None = None
    monitoring_period_end: str | None = None
    verification_date: str | None = None
    issuance_date: str | None = None
    buffer_contribution: float | None = None
    vvb_name: str | None = None
    registry_status: str = "planned"
    notes: str | None = None


@router.post("/{project_id}/issuances")
def add_issuance_endpoint(project_id: int, data: IssuanceCreate):
    from carbongpt.core.lifecycle_manager import add_issuance
    try:
        result = add_issuance(
            project_id, data.vintage_year, data.credits_requested,
            data.credits_issued, data.monitoring_period_start,
            data.monitoring_period_end, data.verification_date,
            data.issuance_date, data.buffer_contribution,
            data.vvb_name, data.registry_status, data.notes,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return result


@router.get("/portfolio/summary")
def get_portfolio_summary_endpoint():
    from carbongpt.core.lifecycle_manager import get_portfolio_summary
    try:
        summary = get_portfolio_summary()
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    return summary
