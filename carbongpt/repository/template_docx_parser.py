"""
template_docx_parser.py — Analyse structurelle d'un template .docx officiel
en champs types et positionnés (docs/SPEC-05.md, T3).

Ne remplace ni ne modifie carbongpt/guides/*.py ni carbongpt/core/
doc_exporter.py — produit une structure séparée (template_fields, T1) que
rien d'autre ne lit encore. L'adaptateur qui les branche à
generate_full_document()/doc_exporter.py (load_guide(), T6) est hors
périmètre de cette session.

Règles de classification (rappelées de docs/SPEC-05.md T3), volontairement
génériques plutôt qu'adaptées table par table — l'objectif de cette session
est un analyseur qui fonctionne sur n'importe quel template Word de ce
type, pas un parseur sur mesure pour le VPA-DD v3.0 :

- paragraphe de style Heading -> field_type='prose', une entrée par
  paragraphe de titre (le contenu narratif qui suit n'est pas segmenté
  séparément, il est réputé appartenir à la section du titre précédent —
  même logique que les guides Python existants, qui associent un bloc de
  prose à un identifiant de section, pas à chaque phrase)
- tableau dont la première cellule commence par "data/parameter" ou
  "data / parameter" (insensible à la casse) -> field_type='parameter_block'
- tableau contenant une case à cocher Word (form field w:ffData/w:checkBox,
  ou glyphe Wingdings de case à cocher — mêmes codes que
  doc_exporter.py::_check_checkbox_in_cell, réutilisés ici pour rester
  cohérent avec ce que l'exporteur sait déjà détecter) -> field_type='checkbox'
- tableau à 2 colonnes où la majorité des lignes ont une deuxième cellule
  vide (motif "étiquette | ___ à remplir") -> field_type='single_value'
- tout autre tableau à plus d'une ligne -> field_type='table'
- Hors périmètre de cette session (T4, session suivante) : détection des
  blocs répétés en checklist_item (motif de code hiérarchique répété type
  "P.1", "P.1.1.1" sur des dizaines de lignes — le cas Safeguarding) et des
  pièces jointes (field_type='attachment'). Un grand tableau de ce type
  est aujourd'hui classé 'table' comme les autres, pas perdu — juste pas
  encore décomposé ligne par ligne.

position toujours enregistrée en indices structurels du document
(paragraph_index pour un titre, table_index pour un tableau) — jamais une
heuristique de texte, condition posée dans SPEC-05.
"""

import logging
import re
from typing import Any

import docx

logger = logging.getLogger(__name__)

_PARAM_BLOCK_HEADER_RE = re.compile(r"^data\s*/\s*parameter\b", re.I)

_CHECKBOX_SYM_CODES = {"00A8", "F0A8", "2610"}
_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_W14_NS = "{http://schemas.microsoft.com/office/word/2010/wordml}"


class TemplateParseError(Exception):
    """Raised when a .docx file cannot be opened or yields zero fields —
    never a silent empty result (same discipline as gs_ingest.py/
    gs_template_ingest.py)."""


def _table_has_checkbox_markup(table) -> bool:
    """Two independent checkbox mechanisms, both checked:

    1. Legacy Word form fields (w:ffData/w:checkBox, or the Wingdings glyph
       codes) — the same signal doc_exporter.py::_check_checkbox_in_cell
       already looks for when it TICKS a checkbox.
    2. Modern structured-document-tag content controls (w:sdt containing a
       w14:checkbox) — confirmed present in VPA-DD v3.0 (479 instances,
       03.08.2026) and ABSENT from doc_exporter.py's detector entirely.
       doc_exporter._check_checkbox_in_cell cannot tick these — it only
       matches legacy form fields. Flagged in docs/STATUS.md for whoever
       picks up T6 (the export side needs its own fix, out of scope here)."""
    for row in table.rows:
        for cell in row.cells:
            if cell._tc.findall(f".//{_W_NS}sdt//{_W14_NS}checkbox"):
                return True
            for p in cell.paragraphs:
                p_el = p._p
                if p_el.findall(f".//{_W_NS}ffData/{_W_NS}checkBox"):
                    return True
                for sym in p_el.findall(f".//{_W_NS}sym"):
                    if sym.get(f"{_W_NS}char") in _CHECKBOX_SYM_CODES:
                        return True
    return False


def _classify_table(table) -> str:
    if not table.rows:
        return "table"

    first_cell_text = table.rows[0].cells[0].text.strip()
    if _PARAM_BLOCK_HEADER_RE.match(first_cell_text):
        return "parameter_block"

    if _table_has_checkbox_markup(table):
        return "checkbox"

    if len(table.columns) == 2 and len(table.rows) > 1:
        empty_second_cell = sum(1 for r in table.rows if not r.cells[1].text.strip())
        if empty_second_cell / len(table.rows) > 0.5:
            return "single_value"

    return "table"


def _table_title(table) -> str:
    """Best-effort label for a table: its first cell's text, truncated —
    not a guessed section title (that would be a text heuristic, which
    SPEC-05 T3 explicitly says position/type must not rely on). Purely
    descriptive, for human review of the extracted structure."""
    if not table.rows:
        return ""
    return table.rows[0].cells[0].text.strip()[:200]


def parse_template_structure(local_path: str) -> list[dict[str, Any]]:
    """
    Parses a .docx template into a flat list of field dicts, each shaped
    for direct insertion into template_fields (minus template_version_id,
    added by the caller):

      {"field_key": str, "parent_section": str | None, "title": str,
       "field_type": str, "position": dict}

    field_key for a heading is a stable slug derived from its position
    ("H27" for the paragraph at index 27) — v3.0 dropped the old "A.1"/"B.1"
    section-code scheme entirely in favour of plain descriptive Word
    headings (verified 03.08.2026, see docs/STATUS.md), so there is no
    reusable natural key to extract from the text itself; the position is
    the only stable identifier available. field_key for a table is
    "T{table_index}".

    Raises TemplateParseError if the file cannot be opened or yields zero
    fields — never an empty result silently accepted.
    """
    try:
        doc = docx.Document(local_path)
    except Exception as exc:
        raise TemplateParseError(f"Could not open {local_path!r} as a .docx: {exc}") from exc

    fields: list[dict[str, Any]] = []
    current_section: str | None = None

    # doc.paragraphs and doc.tables are two SEPARATE flat lists in python-docx
    # — they don't preserve the true interleaved order of the document body.
    # A table's parent_section must be "whichever Heading 1 came before it in
    # the actual document", so the body's real XML children are walked here,
    # each w:p/w:tbl matched back to its python-docx object via a running
    # index — that index is what doc.paragraphs[i]/doc.tables[i] already use,
    # so field_key stays "H{paragraph_index}"/"T{table_index}" as documented.
    para_idx = -1
    table_idx = -1
    for child in doc.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            para_idx += 1
            p = doc.paragraphs[para_idx]
            style = p.style.name if p.style else ""
            text = p.text.strip()
            if not text or not style.startswith("Heading"):
                continue
            level_match = re.match(r"Heading (\d+)", style)
            level = int(level_match.group(1)) if level_match else None
            if level == 1:
                current_section = text
            fields.append({
                "field_key": f"H{para_idx}",
                "parent_section": current_section if level != 1 else None,
                "title": text,
                "field_type": "prose",
                "position": {"paragraph_index": para_idx, "heading_level": level},
            })
        elif tag == "tbl":
            table_idx += 1
            table = doc.tables[table_idx]
            fields.append({
                "field_key": f"T{table_idx}",
                "parent_section": current_section,
                "title": _table_title(table),
                "field_type": _classify_table(table),
                "position": {"table_index": table_idx, "rows": len(table.rows), "columns": len(table.columns)},
            })

    if not fields:
        raise TemplateParseError(
            f"{local_path!r} parsed to zero fields — no headings and no tables found, "
            "refusing to record an empty structure (structure likely unrecognised)"
        )
    return fields
