import io
import os
import re
import copy
import logging
from datetime import datetime

logger = logging.getLogger(__name__)

# Modern structured-document-tag checkboxes (Word 2010+ content controls) —
# see _check_checkbox_in_cell(). Same namespace docs/SPEC-05.md's
# template_docx_parser.py already uses to DETECT these; here they're TICKED.
_W14_NS = "{http://schemas.microsoft.com/office/word/2010/wordml}"

STANDARD_LABELS = {
    "GoldStandard": "Gold Standard",
    "Verra": "Verra VCS",
}

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "document_repository")

TEMPLATE_FILES = {
    ("GoldStandard", "pdd"): "2f6cc71edf6a403a8b4fa2dbbf9877f3_T-PreReview_V1.5-Project-Design-Document.docx",
    ("GoldStandard", "mr"): "72af536a826648149d42683b19cb630d_T-PerfCert_V1.1-Monitoring-Report.docx",
    ("GoldStandard", "vpa_dd"): "1480b5e205a54ed8b6b7b8e86af19ad9_T-PreReview_v2.3-VPA-Design-Document.docx",
    ("GoldStandard", "poa_dd"): "c58aff7a21f044a2bed1cfb5a8e068f5_T-PreReview_V2.2_POA-Design-Document.docx",
    ("Verra", "pdd"): "ae547f753ee34b55927782d5243c4737_VCS-Project-Description-Template-v4.4-FINAL2_1772118270885.docx",
    ("Verra", "mr"): "a7eeb5a6d7264480ac6292ebbb308929_VCS-Monitoring-Report-Template-v4.4-FINAL2.docx",
    ("Verra", "valver"): "344a508a729f444bb6168c6b32124d21_VCS-Joint-Validation-and-Verification-Template-v4.4.docx",
}

GS_SECTION_RE = re.compile(
    r"^([A-F]\.\d+(?:\.\d+)?)\b", re.IGNORECASE
)
VCS_SECTION_RE = re.compile(
    r"^(\d+\.\d+(?:\.\d+)?)\b"
)


def _resolve_template_path(standard, doc_type):
    fname = TEMPLATE_FILES.get((standard, doc_type))
    if not fname:
        return None
    path = os.path.join(TEMPLATE_DIR, fname)
    if os.path.isfile(path):
        return path
    return None


def _parse_markdown_content(content_text):
    lines = content_text.split("\n")
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("|") and i + 1 < len(lines):
            table_lines = [stripped]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("|"):
                table_lines.append(lines[j].strip())
                j += 1
            if len(table_lines) >= 2:
                blocks.append(("table", table_lines))
                i = j
                continue

        if stripped.startswith("#### "):
            blocks.append(("heading4", stripped[5:].strip()))
        elif stripped.startswith("### "):
            blocks.append(("heading4", stripped[4:].strip()))
        elif stripped.startswith("## "):
            blocks.append(("heading3", stripped[3:].strip()))
        elif stripped.startswith("# "):
            blocks.append(("heading3", stripped[2:].strip()))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            blocks.append(("bullet", stripped[2:].strip()))
        elif re.match(r"^\d+\.\s", stripped):
            blocks.append(("numbered", re.sub(r"^\d+\.\s", "", stripped).strip()))
        elif stripped.startswith("[INSERT:") or stripped.startswith("[insert:"):
            blocks.append(("insert_placeholder", stripped))
        elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            blocks.append(("bold_line", stripped[2:-2].strip()))
        elif stripped == "---":
            pass
        elif stripped == "":
            blocks.append(("empty", ""))
        else:
            blocks.append(("paragraph", stripped))
        i += 1
    return blocks


def _parse_md_table(table_lines):
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.split("|")]
        cells = [c for c in cells if c != ""]
        if cells and all(re.match(r"^[-:]+$", c) for c in cells):
            continue
        if cells:
            rows.append(cells)
    return rows


def _strip_leading_title(content_text, section_id):
    lines = content_text.split("\n")
    stripped = []
    skip_next_blank = False
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if i < 3 and line.startswith("#"):
            heading_text = re.sub(r"^#+\s*", "", line).strip()
            sid_pattern = re.compile(
                r"^" + re.escape(section_id) + r"[\s.\-:]+", re.IGNORECASE
            )
            if sid_pattern.match(heading_text) or heading_text.lower().startswith(section_id.lower()):
                skip_next_blank = True
                i += 1
                continue
        if skip_next_blank and line == "":
            skip_next_blank = False
            i += 1
            continue
        skip_next_blank = False
        stripped.append(lines[i])
        i += 1
    return "\n".join(stripped)


def _get_template_font(doc):
    try:
        normal_style = doc.styles["Normal"]
        if normal_style.font.name:
            return normal_style.font.name
    except Exception:
        pass
    return None


def _get_template_font_size(doc):
    try:
        normal_style = doc.styles["Normal"]
        if normal_style.font.size:
            size_half_pts = normal_style.font.size.pt * 2
            return round(size_half_pts) / 2
    except Exception:
        pass
    return None


def _insert_content_into_doc(doc, insert_after_element, content_text, project_info=None, section_id=None):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    if section_id:
        content_text = _strip_leading_title(content_text, section_id)

    template_font = _get_template_font(doc)

    blocks = _parse_markdown_content(content_text)
    parent = insert_after_element.getparent()
    ref_element = insert_after_element

    def _make_rpr(bold=False, italic=False, color=None, font_size=None):
        rpr = doc.element.makeelement(qn("w:rPr"), {})
        if template_font:
            fonts_el = doc.element.makeelement(qn("w:rFonts"), {
                qn("w:ascii"): template_font,
                qn("w:hAnsi"): template_font,
                qn("w:cs"): template_font,
            })
            rpr.append(fonts_el)
        if bold:
            rpr.append(doc.element.makeelement(qn("w:b"), {}))
        if italic:
            rpr.append(doc.element.makeelement(qn("w:i"), {}))
        if color:
            color_el = doc.element.makeelement(qn("w:color"), {qn("w:val"): color})
            rpr.append(color_el)
        if font_size:
            half_pts = str(int(round(font_size * 2)))
            sz_el = doc.element.makeelement(qn("w:sz"), {qn("w:val"): half_pts})
            szcs_el = doc.element.makeelement(qn("w:szCs"), {qn("w:val"): half_pts})
            rpr.append(sz_el)
            rpr.append(szcs_el)
        return rpr

    def _add_para_after(text, bold=False, italic=False, color=None, font_size=None):
        nonlocal ref_element

        p_el = doc.element.makeelement(qn("w:p"), {})
        r_el = doc.element.makeelement(qn("w:r"), {})
        rpr = _make_rpr(bold=bold, italic=italic, color=color, font_size=font_size)

        r_el.append(rpr)
        t_el = doc.element.makeelement(qn("w:t"), {})
        t_el.text = text
        t_el.set(qn("xml:space"), "preserve")
        r_el.append(t_el)
        p_el.append(r_el)

        ref_element.addnext(p_el)
        ref_element = p_el
        return p_el

    def _add_table_after(rows_data):
        nonlocal ref_element
        if not rows_data:
            return

        ncols = max(len(r) for r in rows_data)
        tbl_el = doc.element.makeelement(qn("w:tbl"), {})

        tbl_pr = doc.element.makeelement(qn("w:tblPr"), {})
        tbl_style = doc.element.makeelement(qn("w:tblStyle"), {qn("w:val"): "TableGrid"})
        tbl_pr.append(tbl_style)
        tbl_w = doc.element.makeelement(qn("w:tblW"), {qn("w:w"): "5000", qn("w:type"): "pct"})
        tbl_pr.append(tbl_w)
        tbl_borders = doc.element.makeelement(qn("w:tblBorders"), {})
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            b = doc.element.makeelement(qn(f"w:{border_name}"), {
                qn("w:val"): "single", qn("w:sz"): "4",
                qn("w:space"): "0", qn("w:color"): "000000"
            })
            tbl_borders.append(b)
        tbl_pr.append(tbl_borders)
        tbl_el.append(tbl_pr)

        tbl_grid = doc.element.makeelement(qn("w:tblGrid"), {})
        for _ in range(ncols):
            tbl_grid.append(doc.element.makeelement(qn("w:gridCol"), {}))
        tbl_el.append(tbl_grid)

        for ri, row_cells in enumerate(rows_data):
            tr = doc.element.makeelement(qn("w:tr"), {})
            for ci in range(ncols):
                tc = doc.element.makeelement(qn("w:tc"), {})
                p = doc.element.makeelement(qn("w:p"), {})
                r = doc.element.makeelement(qn("w:r"), {})
                rpr = _make_rpr(bold=(ri == 0), font_size=10)
                r.append(rpr)
                t = doc.element.makeelement(qn("w:t"), {})
                t.text = row_cells[ci] if ci < len(row_cells) else ""
                t.set(qn("xml:space"), "preserve")
                r.append(t)
                p.append(r)
                tc.append(p)
                tr.append(tc)
            tbl_el.append(tr)

        ref_element.addnext(tbl_el)
        ref_element = tbl_el

    base_size = _get_template_font_size(doc)

    for block_type, block_data in blocks:
        if block_type == "table":
            rows = _parse_md_table(block_data)
            if rows:
                _add_table_after(rows)
        elif block_type == "heading3":
            _add_para_after(block_data, bold=True, font_size=base_size + 2 if base_size else 13)
        elif block_type == "heading4":
            _add_para_after(block_data, bold=True, font_size=base_size or 11)
        elif block_type == "bullet":
            _add_para_after(f"\u2022  {block_data}", font_size=base_size)
        elif block_type == "numbered":
            _add_para_after(block_data, font_size=base_size)
        elif block_type == "insert_placeholder":
            _add_para_after(block_data, bold=True, color="CC0000", font_size=base_size)
        elif block_type == "bold_line":
            _add_para_after(block_data, bold=True, font_size=base_size)
        elif block_type == "empty":
            _add_para_after("")
        elif block_type == "paragraph":
            _add_para_after(block_data, font_size=base_size)


def _fill_gs_kpi_table(doc, project_info):
    from docx.oxml.ns import qn

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    kpi_table = doc.tables[0] if doc.tables else None
    if not kpi_table:
        return

    intake = project_info.get("intake", {})
    proponent = intake.get("proponent", {})
    overview = intake.get("project_overview", {})

    for row in kpi_table.rows:
        label = row.cells[0].text.strip().lower()
        cell = row.cells[1]

        is_title_field = False
        if "title" in label:
            if any(w in label for w in ["of project", "of programme", "of vpa", "of cpa"]):
                if "corresponding" not in label and "applicable" not in label and "gs id" not in label:
                    is_title_field = True
        if is_title_field:
            cell.text = project_info.get("name", "")
        elif "host country" in label:
            cell.text = project_info.get("country", "")
        elif "methodology" in label and "applied" in label:
            cell.text = project_info.get("methodology", "")
        elif "project developer" in label and "participant" not in label:
            dev_name = proponent.get("organization_name", "") or proponent.get("contact_person", "") or intake.get("developer_name", "")
            if dev_name:
                cell.text = dev_name
        elif "scale of the project" in label:
            scale = overview.get("scale", "") or intake.get("project_scale", "")
            if scale:
                scale_map = {"Micro-scale": "Micro", "Small-scale": "Small", "Large-scale": "Large"}
                scale_val = scale_map.get(scale, scale)
                _check_checkbox_in_cell(cell._element, scale_val, ns)
        elif "activity requirements" in label:
            activity_type = overview.get("activity_type", "") or intake.get("activity_type", "")
            if activity_type:
                _check_checkbox_in_cell(cell._element, activity_type, ns)
        elif "product requirements" in label:
            _check_checkbox_in_cell(cell._element, "GHG Emissions Reduction", ns)


def _check_checkbox_in_cell(cell_element, target_text, ns):
    target_lower = target_text.lower().strip()

    for para_el in cell_element.findall(".//w:p", ns):
        para_text = ""
        for t_el in para_el.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
            para_text += (t_el.text or "")
        para_text_lower = para_text.strip().lower()

        if not para_text_lower:
            continue

        is_match = (
            target_lower in para_text_lower
            or para_text_lower in target_lower
        )

        if not is_match:
            continue

        ff_datas = para_el.findall(".//w:ffData", ns)
        for ff in ff_datas:
            cb = ff.find("w:checkBox", ns)
            if cb is not None:
                for old_default in cb.findall("w:default", ns):
                    cb.remove(old_default)
                for old_checked in cb.findall("w:checked", ns):
                    cb.remove(old_checked)
                from lxml import etree
                checked_el = etree.SubElement(
                    cb,
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}default"
                )
                checked_el.set(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
                    "1"
                )

        for sym_run in para_el.findall(".//w:r", ns):
            for sym in sym_run.findall("w:sym", ns):
                current_char = sym.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}char", "")
                if current_char in ("00A8", "F0A8", "2610"):
                    sym.set(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}char",
                        "00FE"
                    )

        # Modern content-control checkboxes (w:sdt/w14:checkbox) — confirmed
        # 03.08.2026: VPA-DD v3.0 uses 479 of these, none of the legacy
        # mechanisms above. Without this, export to v3.0 ticks nothing.
        # Two things make a checkbox visually "checked" in Word: the state
        # flag (w14:checked/@w14:val) AND the displayed glyph run, which is
        # a literal Unicode character (☐ U+2610 unchecked), not a Wingdings
        # w:sym char code — confirmed by inspecting the real XML, not assumed.
        for sdt_el in para_el.findall(".//w:sdt", ns):
            checkbox_el = sdt_el.find(f".//{_W14_NS}checkbox")
            if checkbox_el is None:
                continue
            checked_el = checkbox_el.find(f"{_W14_NS}checked")
            checked_state_el = checkbox_el.find(f"{_W14_NS}checkedState")
            if checked_el is not None:
                checked_el.set(f"{_W14_NS}val", "1")
            if checked_state_el is None:
                continue
            checked_codepoint = checked_state_el.get(f"{_W14_NS}val")
            if not checked_codepoint:
                continue
            checked_char = chr(int(checked_codepoint, 16))
            sdt_content = sdt_el.find("w:sdtContent", ns)
            if sdt_content is None:
                continue
            for t_el in sdt_content.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                if t_el.text and ord(t_el.text[:1] or "\0") in (0x2610, 0x2612):
                    t_el.text = checked_char


GS_MR_TEMPLATE_REMAP = {
    "A.3": "A.4",
    "A.4": None,
}

GS_MR_SUBSECTION_ALIASES = {
    "SECTION C. DESCRIPTION OF MONITORING SYSTEM": "C.1",
    "SECTION F. SAFEGUARDS REPORTING": "F.1",
    "B.1.1 Forward Action Requests": "B.1",
    "B.2.1": "B.2",
    "B.2.2": "B.2",
    "B.2.3": "B.2",
    "B.2.4": "B.2",
    "B.2.5": "B.2",
    "E.5.1 Explanation": "E.5",
}

GS_PDD_SUBSECTION_ALIASES = {
    "Eligibility of the project under Gold Standard": "A.1.1",
    "Legal ownership of products generated by the project": "A.1.2",
    "B.5.1 Prior Consideration": "B.5.1",
    "B.5.2 Ongoing Financial Need": "B.5.2",
    "B.6.1 Explanation of methodological choices": "B.6.1",
    "B.6.2 Data and parameters fixed ex ante": "B.6.2",
    "B.6.3 Ex ante estimation of SDG Impact": "B.6.3",
    "B.6.4 Summary of ex ante estimates": "B.6.4",
    "B.7.1 Data and parameters to be monitored": "B.7",
    "B.7.2 Sampling plan": "B.7",
    "B.7.3 Other elements of monitoring plan": "B.7",
    "C.1.1 Start date of project": "C.1",
    "C.1.2 Expected operational lifetime of project": "C.1",
    "C.2.1 Start date of crediting period": "C.2",
    "C.2.2 Total length of crediting period": "C.2",
}


def _fill_gs_template(doc, sections_map, project_info, template_type="pdd"):
    paragraphs = list(doc.paragraphs)

    if template_type == "mr":
        subsection_aliases = GS_MR_SUBSECTION_ALIASES
        section_remap = GS_MR_TEMPLATE_REMAP
    else:
        subsection_aliases = GS_PDD_SUBSECTION_ALIASES
        section_remap = {}

    _fill_gs_kpi_table(doc, project_info)

    filled_sids = set()

    for i, para in enumerate(paragraphs):
        text = para.text.strip()

        sid = None
        m = GS_SECTION_RE.match(text)
        if m:
            raw_sid = m.group(1).upper()
            sid = section_remap.get(raw_sid, raw_sid)
        else:
            for alias_prefix, alias_sid in subsection_aliases.items():
                if text.startswith(alias_prefix):
                    sid = alias_sid
                    break

        if not sid:
            continue

        content = sections_map.get(sid, "")
        if not content:
            continue
        if sid in filled_sids:
            j = i + 1
            if j < len(paragraphs) and paragraphs[j].text.strip() == ">>":
                para_el = paragraphs[j]._element
                parent = para_el.getparent()
                if parent is not None:
                    parent.remove(para_el)
            continue
        filled_sids.add(sid)

        j = i + 1
        while j < len(paragraphs):
            next_text = paragraphs[j].text.strip()
            next_style = paragraphs[j].style.name if paragraphs[j].style else ""
            if next_text == ">>" or (next_text == "" and "Block Text" not in next_style and "Heading" not in next_style):
                para_el = paragraphs[j]._element
                parent = para_el.getparent()
                if parent is not None:
                    parent.remove(para_el)
                    j += 1
                    continue
            break

        _insert_content_into_doc(doc, para._element, content, project_info, section_id=sid)

    _remove_remaining_placeholders_gs(doc, sections_map)


def _remove_remaining_placeholders_gs(doc, sections_map):
    for para in list(doc.paragraphs):
        text = para.text.strip()
        if text != ">>":
            continue
        para_el = para._element
        prev = para_el.getprevious()
        if prev is not None:
            prev_text = prev.text or ""
            for r in prev.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                prev_text = r.text or ""
            prev_full = ""
            for r in prev.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                prev_full += (r.text or "")
            prev_full = prev_full.strip()

            m = GS_SECTION_RE.match(prev_full)
            if m:
                sid = m.group(1).upper()
                if sid in sections_map:
                    parent = para_el.getparent()
                    if parent is not None:
                        parent.remove(para_el)
                    continue

            for alias_prefix in [
                "B.5.1", "B.5.2", "B.6.1", "B.6.2", "B.6.3", "B.6.4",
                "B.7.1", "B.7.2", "B.7.3",
                "C.1.1", "C.1.2", "C.2.1", "C.2.2",
                "Eligibility of the project", "Legal ownership",
                "Additionality", "Applicability", "Compliance",
                "Level of accuracy", "Scale of the project",
                "Stakeholder consultation", "Sustainable development",
                "Safeguarding assessment",
            ]:
                if prev_full.startswith(alias_prefix):
                    parent = para_el.getparent()
                    if parent is not None:
                        parent.remove(para_el)
                    break


GS_VPA_DD_TITLE_MAP = {
    "Purpose and general description": "A.1",
    "Eligibility of the VPA under approved PoA": "A.1.1",
    "Legal ownership of products": "A.1.2",
    "Location of VPA": "A.2",
    "Technologies and/or measures": "A.3",
    "Scale of the VPA": "A.4",
    "Funding sources of VPA": "A.5",
    "Reference of approved methodology": "B.1",
    "Applicability of methodology": "B.2",
    "VPA boundary": "B.3",
    "Establishment and description of baseline scenario": "B.4",
    "Demonstration of additionality": "B.5",
    "Prior Consideration": "B.5.1",
    "Ongoing Financial Need": "B.5.2",
    "Sustainable Development Goals": "B.6",
    "Explanation of methodological choices": "B.6.1",
    "Data and parameters fixed ex ante": "B.6.2",
    "Ex ante estimation of SDG Impact": "B.6.3",
    "Summary of ex ante estimates": "B.6.4",
    "Monitoring plan": "B.7",
    "Data and parameters to be monitored": "B.7.1",
    "Sampling plan": "B.7.2",
    "Other elements of monitoring plan": "B.7.3",
    "Duration of project": "C.1",
    "Start date of VPA": "C.1.1",
    "Expected operational lifetime": "C.1.2",
    "Crediting period": "C.2",
    "Start date of crediting period": "C.2.1",
    "Total length of crediting period": "C.2.2",
    "Safeguarding Principles": "D.1",
    "Assessment that project complies with GS4GG Gender": "D.2",
    "Summary of stakeholder mitigation": "E.1",
    "Final continuous input": "E.2",
    "Eligibility and inclusion criteria": "F.1",
}

GS_POA_DD_TITLE_MAP = {
    "Purpose and general description of the PoA": "A.1",
    "Physical/ Geographical boundary": "A.2",
    "Physical/Geographical boundary": "A.2",
    "Technologies/measures": "A.3",
    "Target/Indicator": "A.4",
    "Coordinating/managing entity": "A.5",
    "Funding sources of PoA": "A.6",
    "Management System": "B.1",
    "Application of methodologies": "B.2",
    "Multiple technologies/measures": "B.2.1",
    "Eligibility criteria for inclusion": "B.3",
    "DEMONSTRATION OF ADDITIONALITY": "C.1",
    "Date of first submission": "D.1",
    "Duration of the PoA": "D.2",
    "Summary of stakeholder consultation at PoA": "E.1",
    "Consideration of stakeholder comments": "E.2",
    "Final Continuous Input / Grievance Mechanism": "E.3",
}


def _match_gs_title_to_sid(text, title_map):
    text_lower = text.lower().strip()
    for prefix, sid in title_map.items():
        if prefix.lower() in text_lower:
            return sid
    return None


def _fill_gs_titled_template(doc, sections_map, project_info, title_map):
    paragraphs = list(doc.paragraphs)

    _fill_gs_kpi_table(doc, project_info)

    filled_sids = set()

    for i, para in enumerate(paragraphs):
        style_name = para.style.name if para.style else ""
        if style_name not in ("Section List", "Section List 2nd", "Section Title"):
            continue

        text = para.text.strip()
        if not text or text == ">>":
            continue

        sid = _match_gs_title_to_sid(text, title_map)
        if not sid:
            continue

        if style_name == "Section Title":
            sid_for_section_title = _match_gs_title_to_sid(text, title_map)
            if sid_for_section_title:
                content = sections_map.get(sid_for_section_title, "")
                if content and sid_for_section_title not in filled_sids:
                    filled_sids.add(sid_for_section_title)
                    j = i + 1
                    while j < len(paragraphs):
                        nt = paragraphs[j].text.strip()
                        ns = paragraphs[j].style.name if paragraphs[j].style else ""
                        if nt == ">>" or (nt == "" and ns not in ("Section List", "Section List 2nd", "Section Title")):
                            para_el = paragraphs[j]._element
                            parent = para_el.getparent()
                            if parent is not None:
                                parent.remove(para_el)
                                j += 1
                                continue
                        break
                    _insert_content_into_doc(doc, para._element, content, project_info, section_id=sid_for_section_title)
            continue

        content = sections_map.get(sid, "")
        if not content:
            continue
        if sid in filled_sids:
            j = i + 1
            if j < len(paragraphs) and paragraphs[j].text.strip() == ">>":
                para_el = paragraphs[j]._element
                parent = para_el.getparent()
                if parent is not None:
                    parent.remove(para_el)
            continue
        filled_sids.add(sid)

        j = i + 1
        while j < len(paragraphs):
            next_text = paragraphs[j].text.strip()
            next_style = paragraphs[j].style.name if paragraphs[j].style else ""
            if next_text == ">>" or (
                next_text == ""
                and next_style not in ("Section List", "Section List 2nd", "Section Title", "Heading 3", "Heading 4", "Heading 5")
            ):
                para_el = paragraphs[j]._element
                parent = para_el.getparent()
                if parent is not None:
                    parent.remove(para_el)
                    j += 1
                    continue
            break

        _insert_content_into_doc(doc, para._element, content, project_info, section_id=sid)

    _remove_remaining_placeholders_gs(doc, sections_map)


VCS_PDD_TITLE_MAP = {
    "Summary Description of the Project": "1.1",
    "Audit History": "1.2",
    "Sectoral Scope and Project Type": "1.3",
    "Project Eligibility": "1.4",
    "Project Design": "1.5",
    "Project Proponent": "1.6",
    "Other Entities Involved in the Project": "1.7",
    "Ownership": "1.8",
    "Project Start Date": "1.9",
    "Project Crediting Period": "1.10",
    "Project Scale and Estimated GHG Emission Reductions or Removals": "1.11",
    "Description of the Project Activity": "1.12",
    "Project Location": "1.13",
    "Conditions Prior to Project Initiation": "1.14",
    "Compliance with Laws, Statutes and Other Regulatory Frameworks": "1.15",
    "Double Counting and Participation under Other GHG Programs": "1.16",
    "Double Claiming, Other Forms of Credit, and Scope 3 Emissions": "1.17",
    "Sustainable Development Contributions": "1.18",
    "Additional Information Relevant to the Project": "1.19",
    "Stakeholder Engagement and Consultation": "2.1",
    "Risks to Stakeholders and the Environment": "2.2",
    "Respect for Human Rights and Equity": "2.3",
    "Ecosystem Health": "2.4",
    "Title and Reference of Methodology": "3.1",
    "Applicability of Methodology": "3.2",
    "Project Boundary": "3.3",
    "Baseline Scenario": "3.4",
    "Additionality": "3.5",
    "Methodology Deviations": "3.6",
    "Baseline Emissions": "4.1",
    "Project Emissions": "4.2",
    "Leakage Emissions": "4.3",
    "Estimated GHG Emission Reductions and Carbon Dioxide Removals": "4.4",
    "Data and Parameters Available at Validation": "5.1",
    "Data and Parameters Monitored": "5.2",
    "Monitoring Plan": "5.3",
}

VCS_MR_TITLE_MAP = {
    "Summary Description of the Implementation Status of the Project": "1.1",
    "Audit History": "1.2",
    "Sectoral Scope and Project Type": "1.3",
    "Project Proponent": "1.4",
    "Other Entities Involved in the Project": "1.5",
    "Project Start Date": "1.6",
    "Project Crediting Period": "1.7",
    "Project Location": "1.8",
    "Title and Reference of Methodology": "1.9",
    "Double Counting and Participation under Other GHG Programs": "1.10",
    "No Double Issuance": "1.10.1",
    "Registration in Other GHG Programs": "1.10.2",
    "Double Claiming, Other Forms of Credit, and Scope 3 Emissions": "1.11",
    "No Double Claiming with Emissions Trading Programs or Binding Emission Limits": "1.11.1",
    "No Double Claiming with Other Forms of Environmental Credit": "1.11.2",
    "Supply Chain (Scope 3) Emissions": "1.11.3",
    "Sustainable Development Contributions": "1.12",
    "Commercially Sensitive Information": "1.13",
    "Stakeholder Engagement and Consultation": "2.1",
    "Stakeholder Identification": "2.1.1",
    "Stakeholder Consultation and Ongoing Communication": "2.1.2",
    "Free, Prior, and Informed Consent": "2.1.3",
    "Grievance Redress Procedure": "2.1.4",
    "Risks to Stakeholders and the Environment": "2.2",
    "Respect for Human Rights and Equity": "2.3",
    "Ecosystem Health": "2.4",
    "Implementation Status of the Project Activity": "3.1",
    "Deviations": "3.2",
    "Grouped Projects": "3.3",
    "Baseline Reassessment": "3.4",
    "Data and Parameters Available at Validation": "4.1",
    "Data and Parameters Monitored": "4.2",
    "Monitoring Plan": "4.3",
    "Baseline Emissions": "5.1",
    "Project Emissions": "5.2",
    "Leakage Emissions": "5.3",
    "GHG Emission Reductions and Carbon Dioxide Removals": "5.4",
}

VCS_VALVER_TITLE_MAP = {
    "Objective": "1.1",
    "Scope and Criteria": "1.2",
    "Reasonableness of Assumptions and Level of Assurance": "1.3",
    "Summary Description of the Project": "1.4",
    "Method and Criteria": "2.1",
    "Document Review": "2.2",
    "Interviews": "2.3",
    "Site Visits": "2.4",
    "Resolution of Findings": "2.5",
    "Forward Action Requests": "2.5.1",
    "Project Details": "3.1",
    "Project Activity Instances in Grouped Projects": "3.2",
    "Safeguards": "3.3",
    "Stakeholder Engagement and Consultation": "3.3.1",
    "Risks to Local Stakeholders and the Environment": "3.3.2",
    "Respect for Human Rights and Equity": "3.3.3",
    "Ecosystem Health": "3.3.4",
    "Application of Methodology": "3.4",
    "Title and Reference": "3.4.1",
    "Applicability": "3.4.2",
    "Project Boundary": "3.4.3",
    "Baseline Scenario": "3.4.4",
    "Additionality": "3.4.5",
    "Quantification of GHG Emission Reductions and Carbon Dioxide Removals": "3.4.6",
    "Methodology Deviations": "3.4.7",
    "Monitoring Plan": "3.4.8",
    "Non-Permanence Risk Analysis": "3.5",
    "Project Implementation Status": "4.1",
    "Accuracy of Reduction and Removal Calculations": "4.2",
    "Quality of Evidence to Determine Reductions and Removals": "4.3",
    "Validation and Verification Summary": "5.1",
    "Validation Conclusion": "5.2",
    "Verification conclusion": "5.3",
    "Ex-ante vs Ex-post ERR Comparison": "5.4",
}


def _match_vcs_title_to_sid(text, title_map):
    normalized = text.strip()
    m = VCS_SECTION_RE.match(normalized)
    if m:
        return m.group(1)
    sid = title_map.get(normalized)
    if sid:
        return sid
    text_lower = normalized.lower()
    for title, map_sid in title_map.items():
        if title.lower() == text_lower:
            return map_sid
    if len(text_lower) > 10:
        for title, map_sid in title_map.items():
            if title.lower() in text_lower or text_lower in title.lower():
                return map_sid
    return None


def _remove_vcs_boilerplate(doc):
    from docx.oxml.ns import qn as _qn
    from docx.text.paragraph import Paragraph as _Para

    body = doc.element.body
    to_remove = []
    instruction_kws = [
        "this template is for", "instructions for completing",
        "file name:", "file type:", "title page formatting",
        "general formatting", "general instructions",
        "note: the instructions", "where a section is not applicable",
        "delete all instructions", "logo (optional)",
        "ddmmmyyyy", "'ddmmmyyyy",
    ]
    passed_first_heading = False
    for child in body:
        if child.tag == _qn('w:p'):
            para = _Para(child, doc)
            style = para.style.name if para.style else ""
            text = para.text.strip()
            text_lower = text.lower()
            if style == "Heading 1" and "instruction" not in text_lower:
                passed_first_heading = True
            if not passed_first_heading:
                if any(kw in text_lower for kw in instruction_kws):
                    to_remove.append(child)
                elif style == "List Paragraph" and ("vcs pd" in text_lower or "ddmmmyyyy" in text_lower):
                    to_remove.append(child)
        elif child.tag == _qn('w:sdt'):
            pass

    for el in to_remove:
        body.remove(el)


def _fill_vcs_kpi_table(doc, project_info):
    if not doc.tables:
        return
    kpi_table = doc.tables[0]
    intake = project_info.get("intake", {})
    proponent = intake.get("proponent", {})

    for row in kpi_table.rows:
        label = row.cells[0].text.strip().lower()
        cell = row.cells[1]

        if "project title" in label:
            cell.text = project_info.get("name", "")
        elif "prepared by" in label:
            dev = proponent.get("organization_name", "") or proponent.get("contact_person", "")
            if dev:
                cell.text = dev
        elif "vcs standard version" in label:
            cell.text = "4.4"

        if cell.text and cell.paragraphs:
            for run in cell.paragraphs[0].runs:
                if run.font:
                    run.font.size = None


def _fill_vcs_template(doc, sections_map, project_info):
    from docx.shared import Pt

    doc_type = project_info.get("doc_type", "pdd")
    if doc_type == "mr":
        title_map = VCS_MR_TITLE_MAP
    elif doc_type == "valver":
        title_map = VCS_VALVER_TITLE_MAP
    else:
        title_map = VCS_PDD_TITLE_MAP

    _remove_vcs_boilerplate(doc)
    _fill_vcs_kpi_table(doc, project_info)

    paragraphs = list(doc.paragraphs)

    for para in paragraphs:
        if para.style and para.style.name == "Template Title":
            text = para.text.strip()
            if "Project" in text or "TITLE" in text:
                para.clear()
                para.add_run(project_info.get("name", "Project Title"))

    filled_sids = set()

    for i, para in enumerate(paragraphs):
        style_name = para.style.name if para.style else ""
        if "Heading" not in style_name:
            continue

        text = para.text.strip()
        if not text:
            continue

        sid = _match_vcs_title_to_sid(text, title_map)
        if not sid:
            continue

        content = sections_map.get(sid, "")
        if not content:
            continue
        if sid in filled_sids:
            continue
        filled_sids.add(sid)

        insert_after = para._element
        next_sibling = insert_after.getnext()
        instructions_removed = 0
        while next_sibling is not None and instructions_removed < 30:
            from docx.oxml.ns import qn as _qn
            if next_sibling.tag != _qn('w:p'):
                if next_sibling.tag == _qn('w:tbl'):
                    from docx.oxml.ns import qn as _qn2
                    next_next = next_sibling.getnext()
                    if next_next is not None and next_next.tag == _qn2('w:p'):
                        from docx.text.paragraph import Paragraph as _P2
                        pp = _P2(next_next, doc)
                        ps = pp.style.name if pp.style else ""
                        if ps in ("Instruction", "Bullets"):
                            pass
                break
            from docx.text.paragraph import Paragraph as _Para
            next_para = _Para(next_sibling, doc)
            ns = next_para.style.name if next_para.style else ""
            nt = next_para.text.strip()

            if ns in ("Instruction", "Bullets", "List Paragraph"):
                to_remove = next_sibling
                next_sibling = next_sibling.getnext()
                to_remove.getparent().remove(to_remove)
                instructions_removed += 1
            elif nt == "" and ns in ("Normal", "Body Text"):
                to_remove = next_sibling
                next_sibling = next_sibling.getnext()
                to_remove.getparent().remove(to_remove)
                instructions_removed += 1
            else:
                break

        _insert_content_into_doc(doc, para._element, content, project_info, section_id=sid)

    logger.info("VCS template filled: %d sections matched out of %d available", len(filled_sids), len(sections_map))


def export_template_word(sections_content, project_info, template_type="pdd"):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    standard = project_info.get("standard", "")
    template_path = _resolve_template_path(standard, template_type)

    sections_map = {}
    for sec in sections_content:
        sid = sec.get("section_id", "")
        content = sec.get("content", "")
        if content and content.strip():
            sections_map[sid] = content

    if template_path:
        try:
            doc = Document(template_path)
            if sections_map:
                if standard == "GoldStandard" and template_type == "vpa_dd":
                    _fill_gs_titled_template(doc, sections_map, project_info, GS_VPA_DD_TITLE_MAP)
                elif standard == "GoldStandard" and template_type == "poa_dd":
                    _fill_gs_titled_template(doc, sections_map, project_info, GS_POA_DD_TITLE_MAP)
                elif standard == "GoldStandard":
                    _fill_gs_template(doc, sections_map, project_info, template_type=template_type)
                elif standard == "Verra":
                    _fill_vcs_template(doc, sections_map, project_info)
                else:
                    _fill_gs_template(doc, sections_map, project_info)
            else:
                if standard == "Verra":
                    _fill_vcs_kpi_table(doc, project_info)
                    _remove_vcs_boilerplate(doc)
                elif standard == "GoldStandard":
                    _fill_gs_kpi_table(doc, project_info)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf
        except Exception as e:
            logger.error("Template-based export failed, falling back to generic: %s", e, exc_info=True)

    return _export_generic_word(sections_content, project_info, template_type)


def _export_generic_word(sections_content, project_info, template_type="pdd"):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    std_label = STANDARD_LABELS.get(project_info.get("standard", ""), project_info.get("standard", ""))
    doc_type_labels = {
        "pdd": "Project Design Document",
        "mr": "Monitoring Report",
        "vpa_dd": "VPA Design Document",
        "poa_dd": "PoA Design Document",
        "valver": "Joint Validation and Verification Report",
    }
    doc_type_label = doc_type_labels.get(template_type, template_type.upper())

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(doc_type_label)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run(std_label)
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    info_items = [
        ("Project Name", project_info.get("name", "[Project Name]")),
        ("Standard", std_label),
        ("Methodology", project_info.get("methodology", "[Methodology]")),
        ("Country", project_info.get("country", "[Country]")),
        ("Document Version", "Draft v1.0"),
        ("Date", datetime.now().strftime("%d %B %Y")),
    ]

    table = doc.add_table(rows=len(info_items), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(info_items):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True if row.cells[0].paragraphs[0].runs else False
        row.cells[1].text = value

    doc.add_page_break()

    toc_heading = doc.add_heading("Table of Contents", level=1)
    for sec in sections_content:
        sid = sec.get("section_id", "")
        title = sec.get("title", "")
        toc_para = doc.add_paragraph(f"{sid}  {title}")
        toc_para.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    for sec in sections_content:
        sid = sec.get("section_id", "")
        title = sec.get("title", "")
        content = sec.get("content", "")
        status = sec.get("status", "placeholder")

        heading = doc.add_heading(f"{sid} {title}", level=2)

        if status == "placeholder":
            placeholder_para = doc.add_paragraph()
            run = placeholder_para.add_run("[This section requires input from the project developer]")
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        if content:
            blocks = _parse_markdown_content(content)
            for block_type, block_data in blocks:
                if block_type == "table":
                    rows = _parse_md_table(block_data)
                    if rows:
                        ncols = max(len(r) for r in rows)
                        tbl = doc.add_table(rows=len(rows), cols=ncols)
                        tbl.style = "Table Grid"
                        for ri, row_cells in enumerate(rows):
                            for ci in range(ncols):
                                cell_text = row_cells[ci] if ci < len(row_cells) else ""
                                tbl.rows[ri].cells[ci].text = cell_text
                                if ri == 0:
                                    for run in tbl.rows[ri].cells[ci].paragraphs[0].runs:
                                        run.bold = True
                elif block_type == "heading3":
                    doc.add_heading(block_data, level=3)
                elif block_type == "heading4":
                    doc.add_heading(block_data, level=4)
                elif block_type == "bullet":
                    doc.add_paragraph(block_data, style="List Bullet")
                elif block_type == "insert_placeholder":
                    p = doc.add_paragraph()
                    run = p.add_run(block_data)
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                    run.font.bold = True
                elif block_type == "bold_line":
                    p = doc.add_paragraph()
                    run = p.add_run(block_data)
                    run.bold = True
                elif block_type == "empty":
                    doc.add_paragraph()
                elif block_type == "paragraph":
                    doc.add_paragraph(block_data)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf




def generate_filled_template(project_info, generated_sections, calc_result=None):
    sections_content = []
    for sec in generated_sections:
        sections_content.append({
            "section_id": sec.get("id", sec.get("section_id", "")),
            "title": sec.get("title", ""),
            "content": sec.get("content", ""),
            "status": "generated" if sec.get("content") else "placeholder",
        })

    if calc_result and not calc_result.get("error"):
        from carbongpt.core.calculation_engine import format_calculation_narrative
        calc_narrative = format_calculation_narrative(calc_result)
        calc_section_ids = {"B.6", "B.6.1", "4.1", "4.2", "4.3", "4.4"}
        inserted = False
        for sec in sections_content:
            sid = sec.get("section_id", "")
            if sid in calc_section_ids or "emission" in sec.get("title", "").lower():
                if not sec.get("content"):
                    sec["content"] = calc_narrative
                    sec["status"] = "generated"
                    inserted = True
                    break
        if not inserted:
            sections_content.append({
                "section_id": "CALC",
                "title": "Emission Reduction Calculations",
                "content": calc_narrative,
                "status": "generated",
            })

    doc_type = project_info.get("doc_type", "pdd")
    return export_template_word(sections_content, project_info, template_type=doc_type)
