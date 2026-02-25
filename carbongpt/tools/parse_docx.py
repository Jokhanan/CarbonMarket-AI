"""
parse_docx.py — Extract structured content from a Word (.docx) document.

Two-pass strategy
-----------------
1. **Heading-style pass** — walks paragraphs looking for Heading 1 / Heading 2
   styles.  If more than one real section is found (beyond PREAMBLE), this
   result is returned.

2. **Heuristic fallback** — if only PREAMBLE (or nothing) was detected via
   headings, a second pass scans *all* paragraphs AND table-cell text for
   known GS MR section markers using regex patterns.

The heuristic is specifically designed for Gold Standard Monitoring Reports
where section titles appear as uppercase bold text or inside tables rather
than as true Word Heading styles.

Returned structure
------------------
{
    "sections": {
        "<SECTION NAME>": "<accumulated text>",
        ...
    },
    "debug": {
        "method": "heading_styles" | "heuristic",
        "raw_paragraphs_count": int,
        "section_markers": [{"index": int, "text": str}, ...],
    }
}
"""

import re
from pathlib import Path
from docx import Document as DocxDocument

_HEADING_STYLES = {"Heading 1", "Heading 2"}
_PREAMBLE_KEY = "PREAMBLE"

_SECTION_PATTERNS = [
    re.compile(r"^KEY\s+PROJECT\s+INFORMATION\b", re.IGNORECASE),
    re.compile(r"^SECTION\s+[A-G]\b", re.IGNORECASE),
]


def _is_heading(paragraph) -> bool:
    return paragraph.style.name in _HEADING_STYLES


def _normalize_marker(text: str) -> str:
    """Collapse whitespace, strip trailing punctuation/dashes for consistent keys."""
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _is_section_marker(text: str) -> bool:
    """Return True if *text* matches a known GS MR section pattern."""
    normalized = _normalize_marker(text)
    return any(p.search(normalized) for p in _SECTION_PATTERNS)


def _extract_all_elements(doc) -> list[dict]:
    """
    Walk the document body in XML order and yield every paragraph and
    table-cell paragraph as dicts of ``{"text": str, "is_heading": bool,
    "source": "paragraph" | "table"}``.

    This ensures table text is included in the stream so heuristic
    detection can find section markers inside tables.
    """
    elements: list[dict] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            from docx.text.paragraph import Paragraph
            para = Paragraph(element, doc)
            text = para.text.strip()
            if text:
                elements.append({
                    "text": text,
                    "is_heading": para.style.name in _HEADING_STYLES,
                    "source": "paragraph",
                })

        elif tag == "tbl":
            from docx.table import Table
            table = Table(element, doc)
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if text:
                            elements.append({
                                "text": text,
                                "is_heading": False,
                                "source": "table",
                            })

    return elements


def _heading_style_pass(elements: list[dict]) -> dict[str, str]:
    """Extract sections using Word heading styles only."""
    sections: dict[str, str] = {}
    current = _PREAMBLE_KEY
    lines: list[str] = []

    def _flush():
        text = " ".join(lines).strip()
        if current in sections:
            sections[current] = (sections[current] + " " + text).strip()
        else:
            sections[current] = text
        lines.clear()

    for el in elements:
        if el["is_heading"]:
            _flush()
            current = el["text"]
        else:
            lines.append(el["text"])

    _flush()

    if _PREAMBLE_KEY in sections and not sections[_PREAMBLE_KEY]:
        del sections[_PREAMBLE_KEY]

    return sections


def _heuristic_pass(elements: list[dict]) -> tuple[dict[str, str], list[dict]]:
    """
    Scan all elements for section markers using regex patterns.
    Returns (sections_dict, markers_list).
    """
    sections: dict[str, str] = {}
    current = _PREAMBLE_KEY
    lines: list[str] = []
    markers: list[dict] = []

    def _flush():
        text = " ".join(lines).strip()
        if current in sections:
            sections[current] = (sections[current] + " " + text).strip()
        else:
            sections[current] = text
        lines.clear()

    for idx, el in enumerate(elements):
        text = el["text"]

        if _is_section_marker(text):
            _flush()
            current = _normalize_marker(text)
            markers.append({"index": idx, "text": current})
        else:
            lines.append(text)

    _flush()

    if _PREAMBLE_KEY in sections and not sections[_PREAMBLE_KEY]:
        del sections[_PREAMBLE_KEY]

    return sections, markers


def parse_docx(file_path: str | Path) -> dict:
    """
    Parse a .docx file and return a sections dictionary.

    Uses heading styles first; falls back to heuristic detection if
    only PREAMBLE (or nothing) is found via headings.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    try:
        doc = DocxDocument(str(path))
    except Exception as exc:
        raise ValueError(f"Could not open document '{path}': {exc}") from exc

    elements = _extract_all_elements(doc)

    heading_sections = _heading_style_pass(elements)

    real_heading_keys = [k for k in heading_sections if k != _PREAMBLE_KEY]

    if len(real_heading_keys) >= 2:
        return {
            "sections": heading_sections,
            "debug": {
                "method": "heading_styles",
                "raw_paragraphs_count": len(elements),
                "section_markers": [],
            },
        }

    heuristic_sections, markers = _heuristic_pass(elements)

    heuristic_real_keys = [k for k in heuristic_sections if k != _PREAMBLE_KEY]

    if len(heuristic_real_keys) >= 1:
        return {
            "sections": heuristic_sections,
            "debug": {
                "method": "heuristic",
                "raw_paragraphs_count": len(elements),
                "section_markers": markers,
            },
        }

    return {
        "sections": heading_sections if heading_sections else {_PREAMBLE_KEY: ""},
        "debug": {
            "method": "heading_styles",
            "raw_paragraphs_count": len(elements),
            "section_markers": [],
        },
    }


def debug_sections(file_path: str | Path) -> dict:
    """
    Diagnostic helper used by the /debug/sections endpoint.

    Returns the first 30 raw paragraphs (original + normalized),
    all detected section markers with indices, and final section names.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    doc = DocxDocument(str(path))
    elements = _extract_all_elements(doc)

    first_30 = []
    for idx, el in enumerate(elements[:30]):
        first_30.append({
            "index": idx,
            "original": el["text"],
            "normalized": _normalize_marker(el["text"]).upper(),
            "source": el["source"],
            "is_heading_style": el["is_heading"],
        })

    _, markers = _heuristic_pass(elements)

    result = parse_docx(file_path)
    section_names = list(result["sections"].keys())

    return {
        "raw_paragraphs_first_30": first_30,
        "section_markers": markers,
        "final_section_names": section_names,
        "detection_method": result["debug"]["method"],
        "total_elements": len(elements),
    }
